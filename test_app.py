"""
Delivery Note Generator Module
"""
import tkinter as tk
from tkinter import ttk, messagebox
from datetime import datetime
import sys
import os
import tkinter.filedialog as fd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from common.utils import replace_placeholder_in_paragraph
# Add parent directory to path for imports
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from common.base_generator import BaseGenerator
from common.utils import format_qty, format_price, format_weight


class DeliveryNoteGenerator(BaseGenerator):
    """Delivery Note Generator - for incoming material deliveries"""
    
    def __init__(self, parent):
        super().__init__(parent, "Delivery Note Generator")

        # Set module-specific attributes before calling super().__init__
        self.module_title = "Delivery Note Generator"
        self.export_button_text = "Export Delivery Note"
        self.default_filename = f"delivery_note_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
                
        # Delivery-specific data
        self.customer_info = {}
        self.delivery_date = tk.StringVar(value=datetime.now().strftime("%Y-%m-%d"))
        self.delivery_location = tk.StringVar()
        self.driver_name = tk.StringVar()
        self.vehicle_info = tk.StringVar()
        self.notes = tk.StringVar()

        self.export_template = self.export_template
        self.create_custom_widgets()


    def create_custom_widgets(self):
        """Create delivery note specific widgets"""
        # Create delivery info section
        self.create_delivery_info_section(self)
        self.get_treeview_columns()
        #self.create_control_buttons()
            
    def create_delivery_info_section(self, parent):
        """Create section for delivery-specific information that resizes with the window"""

        # Allow column 0 of main_frame to expand
        self.main_frame.grid_columnconfigure(0, weight=1)

        info_frame = tk.LabelFrame(
            parent, 
            text="Delivery Information", 
            bg="#F0F0F0", 
            font=("Arial", 12, "bold")
        )
        info_frame.grid(row=0, column=0, padx=15, pady=10, sticky="nsew")  # expand in all directions

        # Allow columns in info_frame to expand proportionally
        for col in range(4):
            info_frame.grid_columnconfigure(col, weight=1)

        # Row 0: Customer and Delivery Date
        tk.Label(info_frame, text="Customer:", bg="#F0F0F0").grid(row=0, column=0, sticky="w", padx=5, pady=5)
        self.customer_entry = tk.Entry(info_frame)
        self.customer_entry.grid(row=0, column=1, padx=5, pady=5, sticky="ew")

        tk.Label(info_frame, text="Delivery Date:", bg="#F0F0F0").grid(row=0, column=2, sticky="w", padx=5, pady=5)
        self.date_entry = tk.Entry(info_frame, textvariable=self.delivery_date)
        self.date_entry.grid(row=0, column=3, padx=5, pady=5, sticky="ew")

        # Row 1: Company and Address
        tk.Label(info_frame, text="Company:", bg="#F0F0F0").grid(row=1, column=0, sticky="w", padx=5, pady=5)
        self.company_entry = tk.Entry(info_frame)
        self.company_entry.insert(0, "Al Mayssan Technical Service")
        self.company_entry.grid(row=1, column=1, padx=5, pady=5, sticky="ew")

        tk.Label(info_frame, text="Address:", bg="#F0F0F0").grid(row=1, column=2, sticky="w", padx=5, pady=5)
        self.address_entry = tk.Entry(info_frame)
        self.address_entry.insert(0, "Dammam, Saudi Arabia")
        self.address_entry.grid(row=1, column=3, padx=5, pady=5, sticky="ew")

        # Row 2: Phone & Fax
        tk.Label(info_frame, text="Phone:", bg="#F0F0F0").grid(row=2, column=0, sticky="w", padx=5, pady=5)
        self.phone_entry = tk.Entry(info_frame)
        self.phone_entry.insert(0, "013-1234567")
        self.phone_entry.grid(row=2, column=1, padx=5, pady=5, sticky="ew")

        tk.Label(info_frame, text="Fax:", bg="#F0F0F0").grid(row=2, column=2, sticky="w", padx=5, pady=5)
        self.fax_entry = tk.Entry(info_frame)
        self.fax_entry.insert(0, "013-7654321")
        self.fax_entry.grid(row=2, column=3, padx=5, pady=5, sticky="ew")

        # Row 3: Customer Number & Q_ID
        tk.Label(info_frame, text="Customer Num:", bg="#F0F0F0").grid(row=3, column=0, sticky="w", padx=5, pady=5)
        self.customer_num_entry = tk.Entry(info_frame)
        self.customer_num_entry.grid(row=3, column=1, padx=5, pady=5, sticky="ew")

        tk.Label(info_frame, text="Quotation ID:", bg="#F0F0F0").grid(row=3, column=2, sticky="w", padx=5, pady=5)
        self.qid_entry = tk.Entry(info_frame)
        self.qid_entry.insert(0, "Q-0001")
        self.qid_entry.grid(row=3, column=3, padx=5, pady=5, sticky="ew")

        # Row 4: Project Name & Notes
        tk.Label(info_frame, text="Project Name:", bg="#F0F0F0").grid(row=4, column=0, sticky="w", padx=5, pady=5)
        self.project_entry = tk.Entry(info_frame)
        self.project_entry.grid(row=4, column=1, padx=5, pady=5, sticky="ew")

        tk.Label(info_frame, text="Notes:", bg="#F0F0F0").grid(row=4, column=2, sticky="w", padx=5, pady=5)
        self.notes_entry = tk.Entry(info_frame, textvariable=self.notes)
        self.notes_entry.grid(row=4, column=3, padx=5, pady=5, sticky="ew")


    '''def create_delivery_info_inline(self, parent_frame):
        """Compact delivery info shown on top right next to logo/title"""
        info_frame = tk.LabelFrame(
            parent_frame, 
            text="Delivery Info", 
            bg="#F0F0F0", 
            font=("Arial", 10, "bold"), 
            labelanchor="n"
        )
        info_frame.pack(side="left", anchor="nw", padx=10, pady=10)

        tk.Label(info_frame, text="Company:", bg="#F0F0F0").grid(row=0, column=0, sticky="w", padx=5, pady=5)
        self.company_entry = tk.Entry(info_frame, width=25)
        self.company_entry.insert(0, "Al Mayssan Technical Service")  # Default
        self.company_entry.grid(row=0, column=1, padx=5, pady=5)

        tk.Label(info_frame, text="Project Name:", bg="#F0F0F0").grid(row=0, column=2, sticky="w", padx=5, pady=5)
        self.project_entry = tk.Entry(info_frame, width=25)
        self.project_entry.grid(row=0, column=3, padx=5, pady=5)

        tk.Label(info_frame, text="Delivery Date:", bg="#F0F0F0").grid(row=1, column=0, sticky="w", padx=5, pady=5)
        self.date_entry = tk.Entry(info_frame, textvariable=self.delivery_date, width=15)
        self.date_entry.grid(row=1, column=1, padx=5, pady=5)

        tk.Label(info_frame, text="Address:", bg="#F0F0F0").grid(row=1, column=2, sticky="w", padx=5, pady=5)
        self.address_entry = tk.Entry(info_frame, width=25)
        self.address_entry.insert(0, "Jeddah, Saudi Arabia")
        self.address_entry.grid(row=1, column=3, padx=5, pady=5)

        tk.Label(info_frame, text="Phone:", bg="#F0F0F0").grid(row=2, column=0, sticky="w", padx=5, pady=5)
        self.phone_entry = tk.Entry(info_frame, width=25)
        self.phone_entry.insert(0, "013-1234567")
        self.phone_entry.grid(row=2, column=1, padx=5, pady=5)

        tk.Label(info_frame, text="Fax:", bg="#F0F0F0").grid(row=2, column=2, sticky="w", padx=5, pady=5)
        self.fax_entry = tk.Entry(info_frame, width=25)
        self.fax_entry.insert(0, "013-7654321")
        self.fax_entry.grid(row=2, column=3, padx=5, pady=5)

        tk.Label(info_frame, text="Customer Num:", bg="#F0F0F0").grid(row=3, column=0, sticky="w", padx=5, pady=5)
        self.customer_num_entry = tk.Entry(info_frame, width=25)
        self.customer_num_entry.grid(row=3, column=1, padx=5, pady=5)

        tk.Label(info_frame, text="Customer:", bg="#F0F0F0").grid(row=3, column=2, sticky="w", padx=5, pady=5)
        self.customer_entry = tk.Entry(info_frame, width=25)
        self.customer_entry.grid(row=3, column=3, padx=5, pady=5)

        tk.Label(info_frame, text="Quotation ID:", bg="#F0F0F0").grid(row=4, column=0, sticky="w", padx=5, pady=5)
        self.qid_entry = tk.Entry(info_frame, width=25)
        self.qid_entry.insert(0, "Q-0001")
        self.qid_entry.grid(row=4, column=1, padx=5, pady=5)'''

    def get_treeview_columns(self):
        """Return columns specific to delivery notes"""
        return ("Part Number", "Description", "Qty", "Supplier", "Unit Price", "Weight", "Status")
    
    def format_item_for_tree(self, product):
        """Format product data for delivery note tree display"""
        qty = format_qty(product.get('Qty', 1))
        price = format_price(product.get('Unit Price', 0), self.currency_unit.get())
        weight = format_weight(product.get('Weight', 0))
        
        return (
            product['Part Number'],
            product['Description'],
            qty,
            product.get('Supplier', ''),
            price,
            weight,
            "Pending"  # Default status for delivery
        )
    
    def export_template(self):
        """Export delivery note template with required columns"""
        try:
            export_data = self.get_export_data()

            if not export_data:
                raise ValueError("No data to export. Please add items to the delivery note.")
            
            #template = self.build_template(export_data)

            filepath = fd.askopenfilename(
                filetypes=[("Word Documents", "*.docx")],
                title="Select Word Template"
            )

            if not filepath:
                return
            
            doc = Document(filepath)

            # Mapping of placeholders
            placeholders = {
                "Company": self.company_entry.get().strip(),
                "Address": self.address_entry.get().strip(),
                "Phone_Num": self.phone_entry.get().strip(),
                "Fax": self.fax_entry.get().strip(),
                "Customer": self.customer_entry.get().strip(),
                "Customer_Num": self.customer_num_entry.get().strip(),
                "Q_ID": self.qid_entry.get().strip(),
                "Project_Name": self.project_entry.get().strip(),
                "Date": self.delivery_date.get().strip()
            }

            # Replace placeholders in paragraphs
            for p in doc.paragraphs:
                for key, val in placeholders.items():
                    replace_placeholder_in_paragraph(p, f"{{{key}}}", val)

            # Replace placeholders in header/footer if needed
            for section in doc.sections:
                for h in section.header.paragraphs:
                    replace_placeholder_in_paragraph(h, f"{{{key}}}", val)
                for f in section.footer.paragraphs:
                    for key, val in placeholders.items():
                        replace_placeholder_in_paragraph(f, f"{{{key}}}", val)

            # Find and populate the first table (assuming it's the item table)
            if doc.tables:
                item_table = None

                # Find the correct table based on its header
                for table in doc.tables:
                    headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
                    if headers[:4] == ['no.', 'item', 'description', 'qty']:
                        item_table = table
                        break

                if not item_table:
                    raise ValueError("Could not find the item table in the Word template.")

                # Remove all rows except header (row 0)
                for row in item_table.rows[1:]:
                    tbl = item_table._tbl
                    tbl.remove(row._tr)

                # Now add one row per item
                for i, item in enumerate(export_data, start=1):
                    row_cells = item_table.add_row().cells
                    row_cells[0].text = str(i)
                    row_cells[1].text = str(item['Part Number'])
                    row_cells[2].text = str(item['Description'])
                    row_cells[3].text = str(item['Qty'])


            # Save output
            save_path = fd.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Document", "*.docx")],
                initialfile="delivery_note.docx"
            )
            if save_path:
                doc.save(save_path)
                messagebox.showinfo("Success", "Document exported successfully.")

        except Exception as e:
            messagebox.showerror("Export Failed", f"{e}")

    '''    def build_template(self, data):
            """Generate the filled template as a string"""
            company = "Al Mayssan Technical Service"
            address = "Dammam, Saudi Arabia"
            phone = "013-1234567"
            fax = "013-7654321"
            project = "N/A"
            q_id = "Q-0001"
            customer_num = "0500123456"

            customer = self.customer_entry.get().strip()
            delivery_date = self.delivery_date.get().strip()

            content = f"""{company}           
    {address}                                                                                                         DN0{datetime.now().strftime('%d-%m-%y')}
    {phone}  		     	
    {fax}                                                                                                                    Date:	
    Attn.:	{customer}                                        		{delivery_date}
    Mob  {customer_num}
    Customer Po Ref:
    Our Quotation: {q_id}

    Project: {project}

    No.\tItem\tDescription\tQty
    """

            for i, item in enumerate(data, 1):
                content += f"{i}\t{item['Part Number']}\t{item['Description']}\t{item['Qty']}\n"

            content += """

    •	  We hereby the recipient confirm that we have received the above items in good order and condition
    •	  We hereby the recipient confirm that we have received the above items as listed

    Supplier            Representative 
        Al Mayssan Technical Service	          Recipient Representative

        Name & Signature	         Name & Signature & Date
    """
            return content        '''
    
    def get_export_data(self):
        """Return data formatted for delivery note export"""
        data = []
        
        # Validate delivery information
        customer = self.customer_entry.get().strip()
        if not customer:
            raise ValueError("Customer name is required")
        
        delivery_date = self.delivery_date.get().strip()
        if not delivery_date:
            raise ValueError("Delivery date is required")
        
        # Process each item in the tree
        for row in self.item_tree.get_children():
            vals = self.item_tree.item(row)['values']
            
            # Basic validation
            if not vals[0] or not vals[1]:  # Part Number and Description
                continue
            
            try:
                qty = int(vals[2].split()[0])  # Extract number from "5 pcs"
                price = float(vals[4].split()[0])  # Extract number from "10.50 SAR"
                weight = float(vals[5].split()[0])  # Extract number from "2.500 kg"
            except (ValueError, IndexError):
                continue
            
            # Calculate total
            total_price = qty * price
            total_weight = qty * weight
            
            row_data = {
                'Delivery Note Date': delivery_date,
                'Customer': customer,
                'Delivery Location': self.delivery_location.get(),
                'Driver': self.driver_name.get(),
                'Vehicle': self.vehicle_info.get(),
                'Part Number': vals[0],
                'Description': vals[1],
                'Supplier': vals[3],
                'Qty': qty,
                'Unit Price (SAR)': price,
                'Total Price (SAR)': round(total_price, 2),
                'Unit Weight (kg)': round(weight, 3),
                'Total Weight (kg)': round(total_weight, 3),
                'Status': vals[6] if len(vals) > 6 else 'Pending',
                'Notes': self.notes.get()
            }
            data.append(row_data)
        
        return data
    

    def insert_material_table(self, doc, placeholder_paragraph, export_data):
        """Insert table in place of the {table} placeholder paragraph"""

        # Get parent element of the paragraph
        parent_element = placeholder_paragraph._element
        parent = parent_element.getparent()
        index = parent.index(parent_element)

        # Remove the placeholder paragraph
        parent.remove(parent_element)

        # Create new table
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "No."
        hdr_cells[1].text = "Part Number"
        hdr_cells[2].text = "Description"
        hdr_cells[3].text = "Qty"

        for i, item in enumerate(export_data, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = str(item['Part Number'])
            row_cells[2].text = str(item['Description'])
            row_cells[3].text = str(item['Qty'])

        # Insert the table XML at the correct position
        table_element = table._element
        parent.insert(index, table_element)

    
    def get_column_width(self, col):
        """Get column width for delivery note treeview"""
        width_map = {
            "Part Number": 100,
            "Description": 200,
            "Qty": 60,
            "Supplier": 120,
            "Unit Price": 80,
            "Weight": 80,
            "Status": 80
        }
        return width_map.get(col, 100)


# Test the module independently
if __name__ == "__main__":
    # Create a simple test window
    root = tk.Tk()
    root.withdraw()  # Hide root window
    
    app = DeliveryNoteGenerator(root)
    app.mainloop()

import os
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from PIL import Image, ImageTk
import pandas as pd
from abc import ABC, abstractmethod

from common.excel_handler import get_products, get_product_details
from common.utils import format_qty, format_price, format_weight


class BaseGenerator(tk.Toplevel, ABC):
    """Base class for all generator modules"""
    
    def __init__(self, parent, title="Generator"):
        super().__init__(parent)
        self.parent = parent
        self.title(title)
        self.geometry("1000x800")
        self.configure(bg="#00A651")
        
        # Common attributes
        self.products = get_products()
        self.currency_unit = tk.StringVar(value="SAR")
        self.selected_items = []
        self.combo_display_list = []
        
        # Build display list for combobox
        self.build_combo_display_list()

        self.main_frame = parent
        #self.main_frame.pack(fill="both", expand=True)

        # ✅ Create scrollable canvas + frame here
        '''self.canvas = tk.Canvas(self.main_frame)
        self.scrollbar = ttk.Scrollbar(self.main_frame, orient="horizontal", command=self.canvas.yview)
        self.scrollable_frame = ttk.Frame(self.canvas)

        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all"))
        )

        self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        self.canvas.configure(yscrollcommand=self.scrollbar.set)

        self.canvas.pack(side="left", fill="both", expand=True)
        self.scrollbar.pack(side="right", fill="y")'''
        
        # Create the UI
        self.create_base_widgets()
        #self.create_custom_widgets()  # Override in subclasses
        
        # Focus on this window
        self.focus_set()
        self.grab_set()
    
    def build_combo_display_list(self):
        """Build the display list for the combobox"""
        self.combo_display_list = [
            f"{p['Part Number']} - {p['Description']}" 
            for p in self.products
        ]
    
    def create_base_widgets(self):
        """Create the base UI components common to all generators"""
        
        # Create main canvas and scrollbar for the entire window
        self.main_canvas = tk.Canvas(self, bg="#00A651", highlightthickness=0)
        self.main_scrollbar = ttk.Scrollbar(self, orient="vertical", command=self.main_canvas.yview)
        self.main_canvas.configure(yscrollcommand=self.main_scrollbar.set)

        # Pack the main canvas and scrollbar
        self.main_canvas.pack(side="left", fill="both", expand=True, padx=15, pady=5)
        self.main_scrollbar.pack(side="right", fill="y", pady=15)

        # Create main frame inside the canvas
        self.main_frame = tk.Frame(self.main_canvas, bg="#00A651")
        self.main_canvas_frame = self.main_canvas.create_window((0, 0), window=self.main_frame, anchor="nw")

        # Bind canvas configurations
        self.main_frame.bind("<Configure>", self.on_main_frame_configure)
        self.main_canvas.bind("<Configure>", self.on_main_canvas_configure)
        
        # Title section
        self.create_title_section()
        
        # Search section
        self.create_search_section()
        
        # Details section
        self.create_details_section()
        
        # Item management section
        self.create_item_management_section()
        
        # Items treeview
        self.create_items_treeview()
        
        # Control buttons
        self.create_control_buttons()
    
    '''def create_title_section(self):
        """Create a horizontal row: delivery info on left, logo + title on right"""
        header_frame = tk.Frame(self.main_frame, bg="#00A651")
        header_frame.pack(fill="x", padx=15, pady=(5, 0))

        # --- LEFT: Delivery Info ---
        self.create_delivery_info_inline(header_frame)

        # --- RIGHT: Logo and Title in one row ---
        right_frame = tk.Frame(header_frame, bg="#00A651")
        right_frame.pack(side="right", anchor="ne", padx=10)

        # Inner frame to hold logo and title side-by-side
        logo_title_frame = tk.Frame(right_frame, bg="#00A651")
        logo_title_frame.pack(anchor="e")

        # Logo
        try:
            logo_path = os.path.join("assets", "mts_logo.png")
            if os.path.exists(logo_path):
                logo_image = Image.open(logo_path).resize((80, 80))
                self.logo_photo = ImageTk.PhotoImage(logo_image)
                logo_label = tk.Label(logo_title_frame, image=self.logo_photo, bg="#00A651")
                logo_label.pack(side="left", anchor="center", padx=(0, 10))
        except Exception as e:
            print(f"Error loading logo: {e}")

        # Title (beside logo)
        title_text = getattr(self, 'module_title', 'Generator')
        title = tk.Label(
            logo_title_frame, 
            text=title_text,
            bg="#00A651", 
            fg="white", 
            font=("Arial", 28, "bold")
        )
        title.pack(side="left", anchor="center")'''


    
    def create_search_section(self):
        """Create the search/product selection section"""
        search_frame = tk.Frame(self.main_frame, bg="#00A651")
        search_frame.pack(pady=10)
        
        tk.Label(
            search_frame, 
            text="Search Products (Part Number, Description):",
            bg="#00A651", 
            fg="white", 
            font=("Arial", 12)
        ).pack(side="left", padx=(0, 10))

        #self.search_var = tk.StringVar()
        self.combo_var = tk.StringVar()
        self.combo = ttk.Combobox(
            search_frame, 
            values=self.combo_display_list, 
            textvariable=self.combo_var, 
            width=50
        )
        self.combo.pack(side="left", pady=5)
        self.combo.bind('<KeyRelease>', self.on_keyrelease)
        self.combo.bind('<<ComboboxSelected>>', self.on_item_selected)
        self.combo.bind('<Return>', self.on_enter_pressed)

        ttk.Button(search_frame, text="Upload File", command=self.upload_file).pack(padx=10)
    
    def create_details_section(self):
        """Create the product details section"""
        outer_frame = tk.Frame(self.main_frame, bg="#F0F0F0")
        outer_frame.pack(fill='x', padx=15, pady=15)
        
        # Add canvas and scrollbar
        self.canvas = tk.Canvas(outer_frame, bg="#F0F0F0")
        scrollbar = ttk.Scrollbar(outer_frame, orient="vertical", command=self.canvas.yview)
        
        # Configure the detail frame
        self.detail_frame = tk.Frame(self.canvas, bg="#F0F0F0", relief="sunken", borderwidth=2, width=300)
        self.detail_frame.pack(padx=15, pady=15)
        
        # Configure canvas
        self.canvas.configure(yscrollcommand=scrollbar.set)
        self.canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Create window in canvas for detail_frame
        self.canvas_frame = self.canvas.create_window((0, 0), window=self.detail_frame, anchor="nw")
        
        # Bind canvas configurations
        self.detail_frame.bind("<Configure>", self.on_frame_configure)
        self.canvas.bind("<Configure>", self.on_canvas_configure)
        
        self.detail_labels = {}
    
    def create_item_management_section(self):
        """Create the item management buttons section"""
        append_item_frame = tk.Frame(self.main_frame, bg='#00A651')
        append_item_frame.pack(pady=10)

        # Add item button
        self.add_btn = ttk.Button(
            append_item_frame, 
            text="Add Item", 
            command=self.add_item
        )
        self.add_btn.grid(row=0, column=1, padx=10)
        self.add_btn.config(state='disabled')

        # Remove Selected button
        self.remove_btn = ttk.Button(
            append_item_frame, 
            text="Remove Item", 
            command=self.remove_selected_item
        )
        self.remove_btn.grid(row=0, column=2, padx=10)
        self.remove_btn.config(state='disabled')
    
    def create_items_treeview(self):
        """Create the items treeview - to be customized by subclasses"""
        columns = self.get_treeview_columns()
        
        # Create a frame to hold the Treeview and Scrollbar
        tree_frame = ttk.Frame(self.main_frame)
        tree_frame.pack(fill="both", expand=True, padx=15, pady=15)

        # Create vertical scrollbar
        tree_scrollbar = ttk.Scrollbar(tree_frame, orient="vertical")
        tree_scrollbar.pack(side="right", fill="y")

        # Create Treeview and attach scrollbar
        self.item_tree = ttk.Treeview(
            tree_frame,
            columns=columns,
            show='headings',
            yscrollcommand=tree_scrollbar.set
        )
        self.item_tree.pack(side="left", fill="both", expand=True)

        # Configure scrollbar
        tree_scrollbar.config(command=self.item_tree.yview)

        # Set up headings and column widths
        for col in columns:
            self.item_tree.heading(col, text=col)
            width = self.get_column_width(col)
            self.item_tree.column(col, width=width)

        # Add double-click binding for editing
        self.item_tree.bind('<Double-1>', self.on_double_click)

    def create_control_buttons(self):
        """Create the control buttons section"""
        btn_frame = tk.Frame(self.main_frame, bg="#00A651")
        btn_frame.pack(pady=5, fill='x')

        ttk.Button(btn_frame, text="Reset", command=self.reset_items).grid(row=0, column=1, padx=10)
        
        # Export button text customizable by subclass
        export_text = getattr(self, 'export_button_text', 'Export to Excel')
        ttk.Button(btn_frame, text=export_text, command=self.export_to_excel).grid(row=0, column=2, padx=10)
        
        if hasattr(self, 'export_template'):
            ttk.Button(btn_frame, text="Export as Template", command=self.export_template).grid(row=0, column=3, padx=10)

        ttk.Button(btn_frame, text="Back to Home", command=self.return_home).grid(row=0, column=4, padx=10)

        self.btn_frame = btn_frame
    
    # Abstract methods to be implemented by subclasses
    @abstractmethod
    def create_custom_widgets(self):
        """Create custom widgets specific to the generator type"""
        pass
    
    @abstractmethod
    def get_treeview_columns(self):
        """Return the columns for the treeview"""
        pass
    
    @abstractmethod
    def get_export_data(self):
        """Return data formatted for export"""
        pass
    
    # Common functionality methods
    def get_column_width(self, col):
        """Get column width for treeview"""
        width_map = {
            "Part Number": 100,
            "Description": 200,
            "Qty": 80,
            "Supplier": 100,
            "Unit Price": 100,
            "Weight": 80,
            "Customer": 120,
            "Delivery Date": 100,
            "Location": 120,
            "Notes": 150
        }
        return width_map.get(col, 100)
    
    def on_keyrelease(self, event):
        """Handle key release in search combobox"""
        typed = self.combo_var.get().lower()
        
        # Clear details first
        self.clear_details()
        self.add_btn.config(state='disabled')
        
        if not typed:
            self.combo['values'] = self.combo_display_list
            return

        # Filter products based on user input
        filtered_products = [
            p for p in self.products
            if (typed in str(p['Part Number']).lower() or
                typed in str(p['Description']).lower())
        ]
        filtered_display = [
            f"{p['Part Number']} - {p['Description']}" for p in filtered_products
        ]
        self.combo['values'] = filtered_display

        # REMOVED: Auto-opening dropdown after each keystroke
        # if filtered_display:
        #     self.combo.event_generate('<Down>')

        # Show all matches in the details frame
        if filtered_products:
            for product in filtered_products:
                self.show_details(product)
    
    def on_item_selected(self, event):
        """Handle item selection from combobox"""
        selection = self.combo_var.get()
        if not selection:
            return

        # Extract part number from selection
        part_number = str(selection.split(' - ')[0].strip())
        print(f"Selected raw combobox text: '{selection}'")
        print(f"Extracted part number: '{part_number}'")
        product = next((p for p in self.products if str(p.get('Part Number', '')).strip() == part_number), None)
        
        if product:
            self.show_details(product)
            self.add_btn.config(state='normal')
        else:
            messagebox.showerror("Error", f"Product with part number '{part_number}' not found.")
            self.add_btn.config(state='disabled')
    

    '''def on_item_selected(self, event):
        self.on_enter_pressed(event)'''

    '''def on_item_selected(self, event):
        """When user selects an item from dropdown, show its details."""
        selection = self.combo.get()
        if not selection:
            return

        # Extract part number from selection and convert to string
        part_number = str(selection.split(' - ')[0].strip())
        product = get_product_details(part_number)
        
        if product:
            self.show_details(product)
            self.add_btn.config(state='normal')
        else:
            messagebox.showerror("Error", f"Product with part number '{part_number}' not found.")
            self.add_btn.config(state='disabled')'''

    def clear_details(self):
        """Clear all widgets in detail frame"""
        for child in self.detail_frame.winfo_children():
            child.destroy()
        
        self.detail_labels = {}
        self.current_product = None
    
    def show_details(self, product):
        """Show product details in detail frame"""
        frame = ttk.Frame(self.detail_frame, borderwidth=1, relief="solid", padding=5, width=280)
        frame.pack(fill='x', pady=5, padx=5)

        # Store the product reference
        self.current_product = product
        frame.product = product

        # Show Part Number and Description
        for key in ['Part Number', 'Description']:
            row = ttk.Frame(frame)
            row.pack(anchor='w', fill='x', padx=2, pady=1)

            key_label = ttk.Label(row, text=f"{key}:", font=('Arial', 8, 'bold'))
            key_label.pack(side='left')

            value_label = ttk.Label(row, text=product[key], font=('Arial', 8))
            value_label.pack(side='left')

        # Buttons frame
        btn_frame = ttk.Frame(frame)
        btn_frame.pack(fill='x', pady=5)

        # Select button
        select_btn = ttk.Button(
            btn_frame, 
            text="Select for List",
            command=lambda p=product, f=frame: self.select_for_list(p, f)
        )
        select_btn.pack(side='left', padx=5)

        # Clear button
        clear_btn = ttk.Button(
            btn_frame,
            text="Clear",
            command=lambda f=frame: f.destroy()
        )
        clear_btn.pack(side='right', padx=5)
    
    def select_for_list(self, product, frame):
        """Add selected product to list"""
        self.current_product = product
        self.add_item(frame)
    
    def add_item(self, frame_to_remove=None):
        """Add currently selected product to the items tree"""
        if not hasattr(self, 'current_product') or self.current_product is None:
            messagebox.showwarning("Warning", "No valid product selected.")
            return

        p = self.current_product
        try:
            # Get formatted values
            values = self.format_item_for_tree(p)
            
            # Insert into tree
            self.item_tree.insert("", "end", values=values)
            self.selected_items.append(p)
            self.update_remove_button_state()

            # Clear selection
            self.combo.set('')
            
            if frame_to_remove:
                frame_to_remove.destroy()
            
            self.add_btn.config(state='disabled')

        except ValueError as e:
            messagebox.showerror("Error", f"Invalid data format: {e}")
    
    def format_item_for_tree(self, product):
        """Format product data for tree display - override in subclasses"""
        qty = format_qty(product.get('Qty', 1))
        price = format_price(product.get('Unit Price', 0), self.currency_unit.get())
        weight = format_weight(product.get('Weight', 0))
        
        return (
            product['Part Number'],
            product['Description'],
            qty,
            product.get('Supplier', ''),
            price,
            weight
        )
    
    def remove_selected_item(self):
        """Remove selected items from the tree"""
        all_items = self.item_tree.get_children()
        num_items = len(all_items)

        if num_items == 0:
            messagebox.showwarning("No Items", "There are no items to remove.")
            self.remove_btn.config(state='disabled')
            return

        if num_items == 1:
            # Only one item - remove it immediately
            self.item_tree.delete(all_items[0])
            if self.selected_items:
                self.selected_items.pop(0)
            self.update_remove_button_state()
            return

        # Multiple items - require selection
        selected = self.item_tree.selection()
        if not selected:
            messagebox.showwarning("No Selection", "Please select item(s) to remove.")
            return

        # Remove selected items
        indices_to_remove = sorted([self.item_tree.index(item) for item in selected], reverse=True)

        for item_id in selected:
            self.item_tree.delete(item_id)

        # Remove from selected_items list
        for index in indices_to_remove:
            if index < len(self.selected_items):
                self.selected_items.pop(index)

        self.update_remove_button_state()
    
    def reset_items(self):
        """Reset all items"""
        self.selected_items.clear()
        for row in self.item_tree.get_children():
            self.item_tree.delete(row)
        self.update_remove_button_state()
    
    def update_remove_button_state(self):
        """Update remove button state based on items"""
        if self.item_tree.get_children():
            self.remove_btn.config(state='normal')
        else:
            self.remove_btn.config(state='disabled')
    
    def export_to_excel(self):
        """Export items to Excel - uses get_export_data() from subclass"""
        if not self.item_tree.get_children():
            messagebox.showinfo("No Data", "No items to export.")
            return

        # Get export data from subclass
        try:
            data = self.get_export_data()
            
            if not data:
                messagebox.showinfo("No Valid Data", "No valid rows to export.")
                return

            # File selection
            default_name = getattr(self, 'default_filename', 'export.xlsx')
            file_path = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel Files", "*.xlsx")],
                title=f"Save {getattr(self, 'module_title', 'Export')} As",
                initialfile=default_name
            )

            
            print(f"User selected file path: {file_path}")

            if not file_path:
                print("Export cancelled by user (no file selected).")
                return
            # Check if file exists and ask for overwrite confirmation
            if os.path.exists(file_path):
                confirm = messagebox.askyesno(
                    "Confirm Overwrite",
                    f"The file '{os.path.basename(file_path)}' already exists.\nDo you want to overwrite it?"
                )
                if not confirm:
                    print("User declined to overwrite existing file.")
                    return

            # Proceed with export
            df_new = pd.DataFrame(data)
            if os.path.exists(file_path):
                # Load existing and combine
                existing_df = pd.read_excel(file_path)
                df_combined = pd.concat([existing_df, df_new], ignore_index=True)
                df_combined.to_excel(file_path, index=False)
            else:
                df_new.to_excel(file_path, index=False)

            messagebox.showinfo("Export Successful", f"File saved to:\n{file_path}")
            print(f"Export successful to {file_path}")

        except Exception as e:
            messagebox.showerror("Export Failed", f"Could not save file:\n{e}")
            print(f"Export failed with error: {e}")

    def upload_file(self):
        """Upload product list file"""
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx *.xls")])
        if file_path:
            try:
                df = pd.read_excel(file_path)
                self.products = df.to_dict(orient='records')
                self.build_combo_display_list()
                self.combo['values'] = self.combo_display_list
                messagebox.showinfo("Success", "Product list uploaded successfully.")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to read Excel file:\n{str(e)}")
    
    def on_double_click(self, event):
        """Handle double-click on treeview item to edit"""
        item = self.item_tree.identify_row(event.y)
        column = self.item_tree.identify_column(event.x)

        if not item:
            return
            
        column_id = int(column[1]) - 1
        column_name = self.item_tree["columns"][column_id]
        
        # Only allow editing Qty and unit price columns by default
        if column_name not in ["Qty", "Unit Price"]:
            return
            
        current_value = self.item_tree.item(item)['values'][column_id]
        
        # Create entry widget for editing
        entry = ttk.Entry(self.item_tree)
        entry.insert(0, current_value)
        
        bbox = self.item_tree.bbox(item, column)
        entry.place(x=bbox[0], y=bbox[1], width=bbox[2], height=bbox[3])
        
        def on_enter(e):
            new_value = entry.get().strip()
            try:
                if column_name == "Qty":
                    qty_str = ''.join(filter(str.isdigit, new_value))
                    if not qty_str:
                        messagebox.showerror("Invalid Input", "Quantity must be a positive number")
                        return
                    
                    qty_val = int(qty_str)
                    if qty_val <= 0:
                        messagebox.showerror("Invalid Input", "Quantity must be greater than 0")
                        return
                    
                    new_value = format_qty(qty_val)
                
                values = list(self.item_tree.item(item)['values'])
                values[column_id] = new_value
                self.item_tree.item(item, values=values)
                
            except ValueError:
                messagebox.showerror("Invalid Input", f"Please enter a valid {column_name.lower()}")
            finally:
                entry.destroy()
        
        entry.bind('<Return>', on_enter)
        entry.bind('<Escape>', lambda e: entry.destroy())
        entry.focus_set()
    
    def on_enter_pressed(self, event):
        """Handle Enter key in search combobox - now opens dropdown"""
        # Open the dropdown when Enter is pressed
        self.combo.event_generate('<Button-1>')
    
    def return_home(self):
        """Return to homepage"""
        self.destroy()
        self.parent.deiconify()
    
    # Canvas scroll methods
    def on_frame_configure(self, event=None):
        """Reset the scroll region to encompass the inner frame"""
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))

    def on_canvas_configure(self, event):
        """When canvas is resized, resize the inner frame to match"""
        width = event.width
        self.canvas.itemconfig(self.canvas_frame, width=width)

    def on_main_frame_configure(self, event=None):
        """Reset the scroll region to encompass the main frame"""
        self.main_canvas.configure(scrollregion=self.main_canvas.bbox("all"))

    def on_main_canvas_configure(self, event):
        """When main canvas is resized, resize the main frame to match"""
        width = event.width
        self.main_canvas.itemconfig(self.main_canvas_frame, width=width)