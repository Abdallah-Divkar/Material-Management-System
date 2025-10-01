"""
Delivery Note Generator Module

This module provides functionality for generating delivery notes with proper
title section layout and consistent styling throughout the application.

Author: Material Management System
Version: 2.0
Last Modified: 2024
"""

from pydoc import doc
import tkinter as tk
import pandas as pd
from tkinter import ttk, messagebox, filedialog
from datetime import datetime
import sys
import os
import tempfile
import json
import tkinter.filedialog as fd
from datetime import datetime
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx2pdf import convert
from common.utils import parse_price_from_display, parse_qty_from_display, parse_weight_from_display, replace_placeholder_in_paragraph, replace_placeholders_in_doc, replace_placeholders_in_doc, save_to_json, load_from_json, parse_float_from_string

# Add parent directory to path for imports
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from common.base_generator import BaseGenerator
from common.utils import format_qty, format_price, format_weight

class DeliveryNoteGenerator(BaseGenerator):
    """
    Delivery Note Generator - for incoming material deliveries
    
    This class extends BaseGenerator to provide delivery note specific functionality
    with improved title section layout and consistent UI structure.
    """
    
    def __init__(self, parent):
        """Initialize the Delivery Note Generator with proper title configuration."""
        # Set module-specific attributes before calling super().__init__
        self.module_title = "Delivery Note Generator"
        self.export_button_text = "Export Delivery Note"
        self.DELIVERY_INFO_FILE = "./backup/client_info_cache.json"
        self.EXCEL_BACKUP_FILE = "./backup/client_info_backup.xlsx"
        
        #self.default_filename = f"delivery_note_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        
        # Initialize parent class
        super().__init__(parent, "Delivery Note Generator")

        #self.saved_delivery_info = {}  # Initialize storage
        # Initialize delivery-specific data
        # Set system date automatically
        #customer_entry = customer_entry.get()
        self.delivery_date_var = tk.StringVar(value=datetime.now().strftime("%d-%m-%y"))
        self.notes = tk.StringVar()
        
        # Create the title section with delivery info callback
        self.create_title_section(left_frame_callback=self.create_delivery_info_inline)
        self.generate_delivery_note_number()
        #self.load_delivery_info()


        # Create custom widgets
        self.create_custom_widgets()

    def create_custom_widgets(self):
        """Create delivery note specific widgets."""
        # This method ensures treeview columns are properly configured
        self.get_treeview_columns()
        '''save_btn = tk.Button(self.btn_frame, text="Save Delivery Info", command=self.save_delivery_note)
        save_btn.grid(row=0, column=5, sticky="w", padx=10, pady=5)

        load_btn = tk.Button(self.btn_frame, text="Load Delivery Info", command=self.load_all_delivery_notes)
        load_btn.grid(row=0, column=6, sticky="w", padx=10, pady=5)'''

    def create_title_section(self, left_frame_callback=None):
        """
        Create an improved title section with proper layout structure.
        
        Args:
            left_frame_callback: Optional callback for creating left frame content
        """
        # Create main header frame with consistent styling
        header_frame = tk.Frame(self.main_frame, bg="#00A695", height=120)
        header_frame.grid(row=0, column=0, sticky="ew", padx=15, pady=(5, 0))
        
        # Configure column weights for responsive layout
        header_frame.grid_columnconfigure(0, weight=1)
        header_frame.grid_columnconfigure(1, weight=0)
        
        # Store header frame reference for child classes
        self.header_frame = header_frame
        
        # Left side: Delivery Information (if callback provided)
        if left_frame_callback:
            self.delivery_info_frame = left_frame_callback(header_frame)
        else:
            # Default empty frame if no callback provided
            self.delivery_info_frame = tk.Frame(header_frame, bg="#F0F0F0")
            self.delivery_info_frame.grid(row=0, column=0, sticky="nsew", padx=(0, 10))
        
        # Right side: Logo and Title section
        self.create_logo_title_section(header_frame)

    def create_logo_title_section(self, parent_frame):
        """
        Create the logo and title section with consistent styling.
        
        Args:
            parent_frame: Parent frame to contain logo and title elements
        """
        logo_title_frame = tk.Frame(parent_frame, bg="#00A695")
        logo_title_frame.grid(row=0, column=1, sticky="ne", padx=(10, 0), pady=10)
        
        # Logo section
        try:
            logo_path = os.path.join("assets", "images",  "mts_logo.png")
            if os.path.exists(logo_path):
                from PIL import Image, ImageTk
                logo_image = Image.open(logo_path).resize((80, 80), Image.Resampling.LANCZOS)
                self.logo_photo = ImageTk.PhotoImage(logo_image)
                logo_label = tk.Label(
                    logo_title_frame, 
                    image=self.logo_photo, 
                    bg="#00A695",
                    relief="flat"
                )
                logo_label.pack(side="left", padx=(0, 15), pady=5)
        except Exception as e:
            print(f"Warning: Could not load logo - {e}")
            # Create placeholder for logo space
            placeholder = tk.Frame(logo_title_frame, width=80, height=80, bg="#00A695")
            placeholder.pack(side="left", padx=(0, 15), pady=5)
        
        # Title section
        title_text = getattr(self, 'module_title', 'Generator')
        title_label = tk.Label(
            logo_title_frame,
            text=title_text,
            bg="#00A695",
            fg="white",
            font=("Arial", 24, "bold"),
            anchor="e",
            justify="right"
        )
        title_label.pack(side="left", pady=5)

    def create_delivery_info_inline(self, parent_frame):
        """
        Create compact delivery information section for the header
        with updated field names and empty entries.
        
        Args:
            parent_frame: Parent frame to contain delivery info
            
        Returns:
            tk.LabelFrame: The created delivery info frame
        """
        info_frame = tk.LabelFrame(
            parent_frame,
            text="Delivery Information",
            bg="#F0F0F0",
            font=("Arial", 11, "bold"),
            labelanchor="nw",
            relief="raised",
            bd=2
        )
        info_frame.grid(row=0, column=0, sticky="nsew", padx=(0, 10), pady=5)
        
        # Configure grid weights for responsive layout
        for i in range(4):
            info_frame.grid_columnconfigure(i, weight=1)

        self.create_info_field(info_frame, "Select Client:", 0, 0)
        self.client_var = tk.StringVar()
        self.client_dropdown = ttk.Combobox(info_frame, textvariable=self.client_var, width=18)
        self.client_dropdown.grid(row=0, column=1)

        # Populate with unique clients
        self.client_dropdown['values'] = self.get_unique_customers()

        # Bind selection to load first delivery note for this client
        self.client_dropdown.bind("<<ComboboxSelected>>", self.on_client_selected)

        self.create_info_field(info_frame, "Delivery Note No.:", 0, 2)
        self.delivery_no_var = tk.StringVar()
        self.delivery_no_entry = self.create_info_entry(
            info_frame, "", 0, 3, textvariable=self.delivery_no_var, width=18
        )

        # Row 0: Customer/Company and Delivery Note Number / Ref
        self.create_info_field(info_frame, "Customer:", 1, 0)
        self.customer_entry = self.create_info_entry(info_frame, "", 1, 1)
        
        #self.create_info_field(info_frame, "Delivery Note Number / Ref:", 0, 2)
        #self.delivery_ref_entry = self.create_info_entry(info_frame, "", 0, 3)
        
        self.create_info_field(info_frame, "Project:", 1, 2)
        self.project_entry = self.create_info_entry(info_frame, "", 1, 3)
        # Row 1: Address
        self.create_info_field(info_frame, "Address:", 2, 0)
        self.address_entry = self.create_info_entry(info_frame, "", 2, 1)

        # Row 2: Phone and Fax
        self.create_info_field(info_frame, "Phone Number:", 2, 2)
        self.phone_entry = self.create_info_entry(info_frame, "", 2, 3)

        # Row 3: Attn. / Contact Person
        self.create_info_field(info_frame, "Attn.:", 3, 0)
        self.incharge_entry = self.create_info_entry(info_frame, "", 3, 1)
        self.create_info_field(info_frame, "Contact  Number:", 3, 2)
        self.contact_number_entry = self.create_info_entry(info_frame, "", 3, 3)
        
        # Row 4: Customer PO Ref and Quotation
        self.create_info_field(info_frame, "Customer PO Ref:", 4, 0)
        self.po_ref_entry = self.create_info_entry(info_frame, "", 4, 1)
        
        self.create_info_field(info_frame, "Quotation:", 4, 2)
        self.quotation_entry = self.create_info_entry(info_frame, "", 4, 3)

        # Row 5: Subject
        self.create_info_field(info_frame, "Subject:", 5, 0)
        self.subject_entry = self.create_info_entry(info_frame, "", 5, 1)
        self.create_info_field(info_frame, "Delivery Date:", 5, 2)
        self.delivery_date = self.create_info_entry(info_frame, self.delivery_date_var.get(), 5, 3, textvariable=self.delivery_date_var)
        
        return info_frame

    def save_delivery_note(self, new_entry=None):
        """
        Save a new delivery note to JSON without overwriting existing notes.
        """
        if new_entry is None:
            new_entry = {
                "Delivery Note No.": self.delivery_no_var.get().strip(),
                "Customer": self.customer_entry.get().strip(),
                "Project": self.project_entry.get().strip(),
                "Address": self.address_entry.get().strip(),
                "Phone": self.phone_entry.get().strip(),
                "Incharge": self.incharge_entry.get().strip(),
                "Customer PO Ref": self.po_ref_entry.get().strip(),
                "Quotation": self.quotation_entry.get().strip(),
                "Subject": self.subject_entry.get().strip(),
                "Contact Number": self.contact_number_entry.get().strip(),
                "Delivery Note Date": self.delivery_date.get().strip(),
                "Notes": self.notes.get().strip()
            }
        filename = self.DELIVERY_INFO_FILE
        # Load existing data first
        all_data = load_from_json(filename)
        
        # Ensure it's a list
        if not isinstance(all_data, list):
            all_data = []

        # Append the new entry
        all_data.append(new_entry)

        # Save everything back
        try:
            with open(filename, "w", encoding="utf-8") as f:
                json.dump(all_data, f, indent=4, ensure_ascii=False)
            messagebox.showinfo("Success", f"Saved delivery note for {new_entry.get('Customer','Unknown')}")
            print(f"[DEBUG] Saved new delivery note for customer '{new_entry.get('Customer', '')}'")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save delivery note:\n{e}")
            print(f"[DEBUG] Error saving JSON: {e}")

        self.append_delivery_note_excel(new_entry)

        # Update client/project cache
        self.update_client_info_cache(
            delivery_no=new_entry["Delivery_Note_No"],
            customer=new_entry["Customer"],
            project=new_entry["Project"],
            address=new_entry.get("Address", ""),
            phone=new_entry.get("Phone", ""),
            incharge=new_entry.get("Incharge", ""),
            contact_number=new_entry.get("Contact Number", ""),
            po_ref=new_entry.get("Customer PO Ref", ""),
            quotation=new_entry.get("Quotation", "")
        )

    def load_all_delivery_notes(self):
        """
        Load all delivery notes safely.
        Ensures return is always a list.
        """
        filename = self.DELIVERY_INFO_FILE
        data = load_from_json(filename)

        if not data:
            print(f"[DEBUG] No delivery notes found in {filename}. Returning empty list.")
            return []

        if not isinstance(data, list):
            print(f"[DEBUG] Warning: Unexpected data format in {filename}. Resetting to [].")
            return []

        print(f"[DEBUG] Loaded {len(data)} delivery notes from {filename}.")
        if len(data) > 0:
            sample = data[0].get("Customer", "Unknown")
            print(f"[DEBUG] Example entry - Customer: {sample}")
        return data

    def get_unique_customers(self, use_cache=False):
        """
        Return a list of unique customers from saved delivery notes.
        """
        if use_cache:
            cache_file = "./backup/client_info_cache.json"
            data = load_from_json(cache_file)
            if not data:
                return []
            customers = list({d.get('Customer', '') for d in data})
        else:
            notes = self.load_all_delivery_notes()
            customers = list({note.get('Customer', '') for note in notes})
            print(f"[DEBUG] Unique customers: {customers}")
        return customers

    def get_notes_by_customer(self, customer):
        """
        Get all delivery notes for a specific customer.
        """
        notes = self.load_all_delivery_notes()
        filtered = [note for note in notes if note.get('Customer', '') == customer]
        print(f"[DEBUG] Found {len(filtered)} notes for customer '{customer}'")
        return filtered

    def create_info_field(self, parent, text, row, col, width=15):
        """
        Create a consistent label field for delivery information.
        
        Args:
            parent: Parent widget
            text: Label text
            row: Grid row
            col: Grid column
            width: Label width
        """
        label = tk.Label(
            parent,
            text=text,
            bg="#F0F0F0",
            font=("Arial", 9, "normal"),
            anchor="w",
            width=width
        )
        label.grid(row=row, column=col, sticky="w", padx=3, pady=0)
        return label

    def create_info_entry(self, parent, default_value="", row=0, col=0, textvariable=None, width=18):
        """
        Create a consistent entry field for delivery information.
        
        Args:
            parent: Parent widget
            default_value: Default entry value
            row: Grid row
            col: Grid column
            textvariable: Optional StringVar for the entry
            width: Entry width
            
        Returns:
            tk.Entry: The created entry widget
        """
        entry = tk.Entry(
            parent,
            font=("Arial", 9),
            width=width,
            relief="sunken",
            bd=1,
            textvariable=textvariable
        )
        if not textvariable and default_value:
            entry.insert(0, default_value)
        entry.grid(row=row, column=col, sticky="ew", padx=0, pady=2)
        return entry

    def load_delivery_note_to_gui(self, note):
        mapping = {
            "Customer": self.customer_entry,
            "Project_Name": self.project_entry,
            "Address": self.address_entry,
            "Phone_Num": self.phone_entry,
            "Incharge": self.incharge_entry,
            "Contact_Num": self.contact_number_entry,
            "Customer_PO": self.po_ref_entry,
            "Quotation": self.quotation_entry,
            "Subject": self.subject_entry,
            "Date": self.delivery_date_var,
            "Delivery_Note_No": self.delivery_no_var
        }
        for key, widget in mapping.items():
            value = note.get(key, "")
            if isinstance(widget, tk.Entry):
                widget.delete(0, tk.END)
                widget.insert(0, value)
            elif isinstance(widget, tk.StringVar):
                widget.set(value)

    def on_client_selected(self, event):
        selected_client = self.client_var.get()
        data = load_from_json("./backup/client_info_cache.json")
        project_info = next((d for d in data if d.get('Customer', '') == selected_client), None)
        if project_info:
            self.customer_entry.delete(0, tk.END)
            self.customer_entry.insert(0, project_info.get("Customer", ""))
            self.project_entry.delete(0, tk.END)
            self.project_entry.insert(0, project_info.get("Project", ""))
            self.address_entry.delete(0, tk.END)
            self.address_entry.insert(0, project_info.get("Address", ""))
            self.phone_entry.delete(0, tk.END)
            self.phone_entry.insert(0, project_info.get("Phone", ""))
            self.incharge_entry.delete(0, tk.END)
            self.incharge_entry.insert(0, project_info.get("Incharge", ""))
            self.contact_number_entry.delete(0, tk.END)
            self.contact_number_entry.insert(0, project_info.get("Contact Number", ""))
            self.po_ref_entry.delete(0, tk.END)
            self.po_ref_entry.insert(0, project_info.get("Customer PO Ref", ""))
            self.quotation_entry.delete(0, tk.END)
            self.quotation_entry.insert(0, project_info.get("Quotation", ""))
            self.subject_entry.delete(0, tk.END)
            self.subject_entry.insert(0, project_info.get("Subject", ""))
        else:
            self.generate_delivery_note_number()


    ''''def on_client_selected(self, event):
        selected_client = self.client_var.get()
        cache_file = "./backup/client_info_cache.json"
        data = load_from_json(cache_file)
        project_info = next((d for d in data if d.get('Customer', '') == selected_client), None)
        if project_info:
            self.customer_entry.delete(0, tk.END)
            self.customer_entry.insert(0, project_info.get("Customer", ""))
            self.project_entry.delete(0, tk.END)
            self.project_entry.insert(0, project_info.get("Project", ""))
            self.address_entry.delete(0, tk.END)
            self.address_entry.insert(0, project_info.get("Address", ""))
            self.phone_entry.delete(0, tk.END)
            self.phone_entry.insert(0, project_info.get("Phone", ""))
            self.incharge_entry.delete(0, tk.END)
            self.incharge_entry.insert(0, project_info.get("Incharge", ""))
        else:
            self.generate_delivery_note_number()'''

    def append_delivery_note_excel(self, new_entry):
        """
        Append a single delivery note to an Excel backup without overwriting previous entries.
        """
        backup_folder = os.path.join(os.getcwd(), "backup")
        os.makedirs(backup_folder, exist_ok=True)
        backup_path = os.path.join(backup_folder, self.EXCEL_BACKUP_FILE)

        df_new = pd.DataFrame([new_entry])

        try:
            # If the backup file exists, append without headers
            if os.path.exists(backup_path):
                df_new.to_excel(backup_path, index=False, header=False, mode='a', engine='openpyxl')
            else:
                # Create a new Excel file with headers
                df_new.to_excel(backup_path, index=False, engine='openpyxl')
            print(f"[DEBUG] Appended delivery note to Excel backup: {backup_path}")
        except Exception as e:
            print(f"[DEBUG] Failed to append Excel backup: {e}")

    def get_treeview_columns(self):
        """Return columns specific to delivery notes."""
        return ("Part Number", "Description", "Qty", "Supplier", "Unit Price", "Weight")
    
    def format_item_for_tree(self, product):
        """
        Format product data for delivery note tree display.
        
        Args:
            product: Product dictionary containing item details
            
        Returns:
            tuple: Formatted values for tree display
        """
        qty = format_qty(product.get('Qty', 1))
        price = format_price(product.get('Unit Price', 0), self.currency_unit.get())
        weight = format_weight(product.get('Weight', 0))
        
        return (
            product['Part Number'],
            product['Description'],
            qty,
            product.get('Supplier', ''),
            price,
            weight
        )
    
    def generate_delivery_note_number(self):
        """
        Generate a new delivery note number in the format DN{seq}-{MM}-{YY}.
        Sequence resets every year.
        """
        all_notes = self.load_all_delivery_notes()
        if not all_notes:
            seq = 1
        else:
            last_dn = all_notes[-1].get("Delivery Note No.", "")
            print(f"[DEBUG] Last delivery note number: {last_dn}")
            try:
                seq = int(last_dn[2:5]) + 1
            except:
                seq = len(all_notes) + 1
        now = datetime.now()
        dn_number = f"DN{seq:03d}-{now.strftime('%m-%y')}"
        self.delivery_no_var.set(dn_number)
        return dn_number
 
    def get_export_data(self):
        """
        Return data formatted for delivery note export.
        """
        data = []

        # Fetch and validate main fields
        customer = self.customer_entry.get().strip()
        if not customer:
            raise ValueError("Customer name is required")

        delivery_date = self.delivery_date.get().strip()
        if not delivery_date:
            raise ValueError("Delivery date is required")

        children = self.item_tree.get_children()
        if not children:
            print("[DEBUG] No items found in treeview!")
            return []

        for idx, row in enumerate(children):
            vals = self.item_tree.item(row)['values']

            # Skip rows missing essential info
            if not vals[0] or not vals[1]:
                print(f"[DEBUG] Skipping row {idx}: missing Part Number or Description")
                continue

            # Parse numeric fields using utility functions
            try:

                qty = parse_qty_from_display(vals[2])
                price = parse_price_from_display(vals[4])
                weight = parse_weight_from_display(vals[5])
            except Exception as e:
                print(f"[DEBUG] Skipping row {idx} due to parse error: {e}")
                continue

            row_data = {
                'Delivery Note No.': self.delivery_no_var.get().strip(),
                'Delivery Note Date': delivery_date,
                'Customer': customer,
                'Project': self.project_entry.get().strip(),
                'Address': self.address_entry.get().strip(),
                'Phone': self.phone_entry.get().strip(),
                'Attn.': self.incharge_entry.get().strip(),
                'Customer PO Ref': self.po_ref_entry.get().strip(),
                'Quotation': self.quotation_entry.get().strip(),
                'Subject': self.subject_entry.get().strip(),
                'Contact Number': self.contact_number_entry.get().strip(),
                'Part Number': vals[0],
                'Description': vals[1],
                'Supplier': vals[3],
                'Qty': qty,
                'Unit Price': round(price, 2),
                'Unit Weight (kg)': round(weight, 3),
                'Total Price': round(qty * price, 2),
                'Total Weight (kg)': round(qty * weight, 3),
                'Notes': self.notes.get().strip()
            }

            data.append(row_data)

        print(f"[DEBUG] Total exportable rows: {len(data)}")
        return data
    
    def export_template(self):
        """
        Export delivery note template with required columns (uses default template).
        """
        try:
            export_data = self.get_export_data()

            if not export_data:
                raise ValueError("No data to export. Please add items to the delivery note.")

            # ✅ Use default template path instead of asking
            default_template = os.path.join(os.getcwd(), "assets", "templates", "delivery_note_template.docx")

            if not os.path.exists(default_template):
                raise FileNotFoundError(f"Default template not found:\n{default_template}")

            doc = Document(default_template)

            # Mapping of placeholders to actual values
            placeholders = {
                "Delivery_Note_No": self.delivery_no_var.get().strip(),
                "Customer": self.customer_entry.get().strip(),
                "Project_Name": self.project_entry.get().strip(),
                "Address": self.address_entry.get().strip(),
                "Phone_Num": self.phone_entry.get().strip(),
                "Incharge": self.incharge_entry.get().strip(),
                "Customer_PO": self.po_ref_entry.get().strip(),
                "Quotation": self.quotation_entry.get().strip(),
                "Subject": self.subject_entry.get().strip(),
                "Contact_Num": self.contact_number_entry.get().strip(),
                "Date": self.delivery_date.get().strip()
            }

            # ✅ Replace in doc (paragraphs, headers, tables…)
            replace_placeholders_in_doc(doc, placeholders)

            # ✅ Populate items
            self.populate_item_table(doc, export_data)

            delivery_no = self.delivery_no_var.get().strip() or "DN"
            customer_name = self.customer_entry.get().strip() or "Customer"
            customer_name = customer_name.replace(" ", "_")
            default_filename = f"{delivery_no}-{customer_name}.docx"

            export_folder = os.path.join(os.getcwd(), "assets", "exports")
            os.makedirs(export_folder, exist_ok=True)

            save_path = fd.asksaveasfilename(
                initialdir=export_folder,
                initialfile=default_filename,
                defaultextension=".docx",
                filetypes=[("Word Document", "*.docx")]
            )
            if not save_path:
                return

            doc.save(save_path)
            save_to_json(export_data)
            messagebox.showinfo("Success", f"Delivery note exported successfully:\n{save_path}")

            self.save_delivery_note()
            print(f"[DEBUG] Delivery info saved to {self.DELIVERY_INFO_FILE}")

        except Exception as e:
            messagebox.showerror("Export Failed", f"Failed to export delivery note:\n{str(e)}")
    
    def print_delivery_note_pdf(self):
        """
        Export delivery note as PDF and send to the system printer.
        """
        try:
            # Step 2a: Export data
            export_data = self.get_export_data()
            if not export_data:
                messagebox.showwarning("No Data", "No items to export.")
                return

            doc_number  = self.delivery_no_var.get().strip() or "DN"
            client_name = self.customer_entry.get().strip() or "Client"
            pdf_filename = f"{doc_number}-{client_name}.pdf"
            temp_dir = tempfile.gettempdir()
            pdf_path = os.path.join(temp_dir, pdf_filename)


            default_template = os.path.join(os.getcwd(), "assets", "templates", "delivery_note_template.docx")

            if not os.path.exists(default_template):
                raise FileNotFoundError(f"Default template not found:\n{default_template}")

            # Fill template with placeholders
            doc = Document(default_template)
            placeholders = {
                "Delivery_Note_No": self.delivery_no_var.get().strip(),
                "Customer": self.customer_entry.get().strip(),
                "Project_Name": self.project_entry.get().strip(),
                "Address": self.address_entry.get().strip(),
                "Phone_Num": self.phone_entry.get().strip(),
                "Incharge": self.incharge_entry.get().strip(),
                "Customer_PO": self.po_ref_entry.get().strip(),
                "Quotation": self.quotation_entry.get().strip(),
                "Subject": self.subject_entry.get().strip(),
                "Contact_Num": self.contact_number_entry.get().strip(),
                "Date": self.delivery_date.get().strip()
            }
            # ✅ Apply everywhere (paragraphs, headers, tables)
            replace_placeholders_in_doc(doc, placeholders)

            self.populate_item_table(doc, export_data)

            temp_docx = os.path.join(temp_dir, "temp_delivery_note.docx")
            doc.save(temp_docx)

            # Convert to PDF
            convert(temp_docx, pdf_path)

            # Print
            if os.name == "nt":  # Windows
                os.startfile(pdf_path, "print")
            elif sys.platform == "darwin":  # macOS
                os.system(f"lp '{pdf_path}'")
            else:  # Linux
                os.system(f"lp '{pdf_path}'")

            messagebox.showinfo("Success", "Delivery note sent to printer.")

            self.save_delivery_note()
            print("[DEBUG] Delivery info saved automatically on print")
        except Exception as e:
            messagebox.showerror("Print Failed", f"Failed to print delivery note:\n{e}")

    def populate_item_table(self, doc, export_data):
        """
        Populate the item table in the Word document.
        
        Args:
            doc: Word document object
            export_data: List of item data dictionaries
        """
        if not doc.tables:
            return

        # Find the correct table (assuming first table with proper headers)
        item_table = None
        for table in doc.tables:
            if len(table.rows) > 0:
                headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
                if 'no.' in headers and 'item' in headers and 'description' in headers:
                    item_table = table
                    break

        if not item_table:
            raise ValueError("Could not find the item table in the Word template.")

        # Clear existing rows (except header)
        for row in item_table.rows[1:]:
            table_element = item_table._tbl
            table_element.remove(row._tr)

        # Add new rows for each item
        for i, item in enumerate(export_data, start=1):
            row_cells = item_table.add_row().cells
            if len(row_cells) >= 4:
                row_cells[0].text = str(i)
                row_cells[1].text = str(item['Part Number'])
                row_cells[2].text = str(item['Description'])
                row_cells[3].text = str(item['Qty'])

    def get_column_width(self, col):
        """
        Get column width for delivery note treeview.
        
        Args:
            col: Column name
            
        Returns:
            int: Column width in pixels
        """
        width_map = {
            "Part Number": 100,
            "Description": 200,
            "Qty": 60,
            "Supplier": 120,
            "Unit Price": 80,
            "Weight": 80
        }
        return width_map.get(col, 100)
    
    def upload_file(self):
        file_data = super().upload_file()
        print(f"[DEBUG] Uploaded file data: {file_data}")
        if file_data:
            self.autofill_from_file_data(file_data)
        else:
            print("[DEBUG] No data extracted from uploaded file.")

    def autofill_from_file_data(self, file_data, doc=None):
        """
        Autofill delivery info entries and TreeView using data from upload_file().
        file_data should be a dict like: {"info": {...}, "items": [...]}
        """
        if not file_data:
            return
        
        info = file_data.get("info", {})
        self.customer_entry.delete(0, tk.END)
        self.customer_entry.insert(0, info.get("Customer", ""))
        self.project_entry.delete(0, tk.END)
        self.project_entry.insert(0, info.get("Project_Name", ""))
        self.address_entry.delete(0, tk.END)
        self.address_entry.insert(0, info.get("Address", ""))
        self.phone_entry.delete(0, tk.END)
        self.phone_entry.insert(0, info.get("Phone_Num", ""))
        self.incharge_entry.delete(0, tk.END)
        self.incharge_entry.insert(0, info.get("Incharge", ""))
        self.contact_number_entry.delete(0, tk.END)
        self.contact_number_entry.insert(0, info.get("Contact_Num", ""))
        self.po_ref_entry.delete(0, tk.END)
        self.po_ref_entry.insert(0, info.get("Customer_PO", ""))
        self.quotation_entry.delete(0, tk.END)
        self.quotation_entry.insert(0, info.get("Quotation", ""))
        self.subject_entry.delete(0, tk.END)
        self.subject_entry.insert(0, info.get("Subject", ""))
        self.delivery_date_var.set(info.get("Date", datetime.now().strftime("%d-%m-%y")))
        self.delivery_no_var.set(info.get("Delivery_Note_No", self.generate_delivery_note_number()))

        self.item_tree.delete(*self.item_tree.get_children())
        for row in file_data.get("items", []):
            self.item_tree.insert("", "end", values=row)
        
        print("DEBUG: autofilling GUI with info:", file_data.get("info"))
        print("DEBUG: items:", file_data.get("items"))


        '''if doc and doc.tables:
            table = doc.tables[0]  # assuming first table has items
            for row in table.rows[1:]:  # skip header
                cells = [cell.text.strip() for cell in row.cells]
                # Parse numeric fields if necessary
                part_no = cells[1]
                desc = cells[2]
                qty = parse_float_from_string(cells[3])
                self.item_tree.insert("", "end", values=(part_no, desc, qty))'''

        messagebox.showinfo("Success", "Word document uploaded and GUI autofilled successfully.") 

# Test the module independently
if __name__ == "__main__":
    # Create a simple test window
    root = tk.Tk()
    root.withdraw()  # Hide root window

# Test the module independently
if __name__ == "__main__":
    # Create a simple test window
    root = tk.Tk()
    root.withdraw()  # Hide root window
    
    try:
        app = DeliveryNoteGenerator(root)
        app.mainloop()
    except Exception as e:
        print(f"Error running delivery note generator: {e}")
        messagebox.showerror("Error", f"Failed to start application:\n{e}")
    finally:
        root.destroy()