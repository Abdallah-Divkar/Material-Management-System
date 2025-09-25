"""
Dispatch Note Generator Module
"""
import json
import tempfile
import tkinter as tk
from tkinter import ttk
from datetime import datetime
import sys
import os
from datetime import datetime
from tkinter import messagebox
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx2pdf import convert
import pandas as pd
from common.utils import replace_placeholder_in_paragraph, save_to_json, load_from_json, parse_float_from_string

# Add parent directory to path for imports
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from common.base_generator import BaseGenerator
from common.utils import format_qty, format_price, format_weight


class DispatchNoteGenerator(BaseGenerator):
    """Dispatch Note Generator - for outgoing material dispatches"""
    
    def __init__(self, parent):
        # Set module-specific attributes before calling super().__init__
        self.module_title = "Dispatch Note Generator"
        self.export_button_text = "Export Dispatch Note"
        self.DISPATCH_INFO_FILE = "./backup/client_info_cache.json"
        self.EXCEL_BACKUP_FILE = "./backup/client_info_backup.xlsx"
        #self.default_filename = f"dispatch_note_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        super().__init__(parent, "Dispatch Note Generator")

        self.dispatch_date_var = tk.StringVar(value=datetime.now().strftime("%d-%m-%y"))
        self.notes = tk.StringVar()

        self.create_title_section(left_frame_callback=self.create_dispatch_info_inline)
        self.create_custom_widgets()
        

    def create_title_section(self, left_frame_callback=None):
        """
        Create an improved title section with proper layout structure.
        
        Args:
            left_frame_callback: Optional callback for creating left frame content
        """
        # Create main header frame with consistent styling
        header_frame = tk.Frame(self.main_frame, bg="#00A695", height=120)
        header_frame.grid(row=0, column=0, sticky="ew", padx=15, pady=(5, 0))
        
        # Configure column weights for responsive layout
        header_frame.grid_columnconfigure(0, weight=1)  # Left side (dispatch info)
        header_frame.grid_columnconfigure(1, weight=0)  # Right side (logo + title)
        
        # Store header frame reference for child classes
        self.header_frame = header_frame
        # Left side: Dispatch Information (if callback provided)
        if left_frame_callback:
            self.dispatch_info_frame = left_frame_callback(header_frame)
        else:
            # Default empty frame if no callback provided
            self.dispatch_info_frame = tk.Frame(header_frame, bg="#F0F0F0")
            self.dispatch_info_frame.grid(row=0, column=0, sticky="nsew", padx=(0, 10))

        # Right side: Logo and Title section
        self.create_logo_title_section(header_frame)

    def create_logo_title_section(self, parent_frame):
        """
        Create the logo and title section with consistent styling.
        
        Args:
            parent_frame: Parent frame to contain logo and title elements
        """
        logo_title_frame = tk.Frame(parent_frame, bg="#00A695")
        logo_title_frame.grid(row=0, column=1, sticky="ne", padx=(10, 0), pady=10)
        
        # Logo section
        try:
            logo_path = os.path.join("assets", "mts_logo.png")
            if os.path.exists(logo_path):
                from PIL import Image, ImageTk
                logo_image = Image.open(logo_path).resize((80, 80), Image.Resampling.LANCZOS)
                self.logo_photo = ImageTk.PhotoImage(logo_image)
                logo_label = tk.Label(
                    logo_title_frame, 
                    image=self.logo_photo, 
                    bg="#00A695",
                    relief="flat"
                )
                logo_label.pack(side="left", padx=(0, 15), pady=5)
        except Exception as e:
            print(f"Warning: Could not load logo - {e}")
            # Create placeholder for logo space
            placeholder = tk.Frame(logo_title_frame, width=80, height=80, bg="#00A695")
            placeholder.pack(side="left", padx=(0, 15), pady=5)
        
        # Title section
        title_text = getattr(self, 'module_title', 'Generator')
        title_label = tk.Label(
            logo_title_frame,
            text=title_text,
            bg="#00A695",
            fg="white",
            font=("Arial", 24, "bold"),
            anchor="e",
            justify="right"
        )
        title_label.pack(side="left", pady=5)
    
    def create_dispatch_info_inline(self, parent_frame):
        """
        Create compact dispatch information section for the header
        with updated field names and empty entries.
        
        Args:
            parent_frame: Parent frame to contain dispatch info
            
        Returns:
            tk.LabelFrame: The created dispatch info frame
        """
        info_frame = tk.LabelFrame(
            parent_frame,
            text="Dispatch Information",
            bg="#F0F0F0",
            font=("Arial", 11, "bold"),
            labelanchor="nw",
            relief="raised",
            bd=2
        )
        info_frame.grid(row=0, column=0, sticky="nsew", padx=(0, 10), pady=5)
        
        # Configure grid weights for responsive layout
        for i in range(4):
            info_frame.grid_columnconfigure(i, weight=1)

        self.create_info_field(info_frame, "Select Client:", 0, 0)
        self.client_var = tk.StringVar()
        self.client_dropdown = ttk.Combobox(info_frame, textvariable=self.client_var, width=20)
        self.client_dropdown.grid(row=0, column=1, padx=5, pady=2)

        # Populate with unique clients
        self.client_dropdown['values'] = self.get_unique_customers()

        # Bind selection to load first dispatch note for this client
        self.client_dropdown.bind("<<ComboboxSelected>>", self.on_client_selected)
        
        # Row 0: Customer/Company 
        self.create_info_field(info_frame, "Client:", 1, 0)
        self.customer_entry = self.create_info_entry(info_frame, "", 1, 1)
        
        # Row 1: Address
        self.create_info_field(info_frame, "Dispatch Address:", 1, 2)
        self.address_entry = self.create_info_entry(info_frame, "", 1, 3)

        # Row 2: Phone and Fax
        self.create_info_field(info_frame, "Phone Number:", 2, 0)
        self.phone_entry = self.create_info_entry(info_frame, "", 2, 1)
        
        self.create_info_field(info_frame, "Fax:", 2, 2)
        self.fax_entry = self.create_info_entry(info_frame, "", 2, 3)
        
        # Row 3: Incharge / Contact Person
        self.create_info_field(info_frame, "Incharge:", 3, 0)
        self.incharge_entry = self.create_info_entry(info_frame, "", 3, 1)
        self.create_info_field(info_frame, "Contact  Number:", 3, 2)
        self.contact_number_entry = self.create_info_entry(info_frame, "", 3, 3)
        
        # Row 4: Customer PO Ref and Quotation
        self.create_info_field(info_frame, "Customer PO Ref:", 4, 0)
        self.po_ref_entry = self.create_info_entry(info_frame, "", 4, 1)
        
        self.create_info_field(info_frame, "Quotation:", 4, 2)
        self.quotation_entry = self.create_info_entry(info_frame, "", 4, 3)
        
        # Row 5: Project
        self.create_info_field(info_frame, "Project:", 5, 0)
        self.project_entry = self.create_info_entry(info_frame, "", 5, 1)
        self.create_info_field(info_frame, "Dispatch Date:", 5, 2)
        self.dispatch_date = self.create_info_entry(info_frame, self.dispatch_date_var.get(), 5, 3, textvariable=self.dispatch_date_var)
        
        return info_frame

    def create_info_field(self, parent, text, row, col, width=15):
        """
        Create a consistent label field for dispatch information.
        
        Args:
            parent: Parent widget
            text: Label text
            row: Grid row
            col: Grid column
            width: Label width
        """
        label = tk.Label(
            parent,
            text=text,
            bg="#F0F0F0",
            font=("Arial", 9, "normal"),
            anchor="w",
            width=width
        )
        label.grid(row=row, column=col, sticky="w", padx=3, pady=2)
        return label

    def create_info_entry(self, parent, default_value="", row=0, col=0, textvariable=None, width=18):
        """
        Create a consistent entry field for dispatch information.
        
        Args:
            parent: Parent widget
            default_value: Default entry value
            row: Grid row
            col: Grid column
            textvariable: Optional StringVar for the entry
            width: Entry width
            
        Returns:
            tk.Entry: The created entry widget
        """
        entry = tk.Entry(
            parent,
            font=("Arial", 9),
            width=width,
            relief="sunken",
            bd=1,
            textvariable=textvariable
        )
        if not textvariable and default_value:
            entry.insert(0, default_value)
        entry.grid(row=row, column=col, sticky="ew", padx=3, pady=2)
        return entry
    
    def create_custom_widgets(self):
        """Create dispatch note specific widgets"""
        self.get_treeview_columns()
        save_btn = tk.Button(self.btn_frame, text="Save Dispatch Info", command=self.save_dispatch_note)
        save_btn.grid(row=0, column=5, sticky="w", padx=10, pady=5)

        load_btn = tk.Button(self.btn_frame, text="Load Dispatch Info", command=self.load_all_dispatch_notes)
        load_btn.grid(row=0, column=6, sticky="w", padx=10, pady=5)

    def save_dispatch_note(self, new_entry=None):
        """
        Save a new dispatch note to JSON without overwriting existing notes.
        """
        if new_entry is None:
            new_entry = {
                "Customer": self.customer_entry.get().strip(),
                "Address": self.address_entry.get().strip(),
                "Phone": self.phone_entry.get().strip(),
                "Fax": self.fax_entry.get().strip(),
                "Incharge": self.incharge_entry.get().strip(),
                "Customer PO Ref": self.po_ref_entry.get().strip(),
                "Quotation": self.quotation_entry.get().strip(),
                "Project": self.project_entry.get().strip(),
                "Contact Number": self.contact_number_entry.get().strip(),
                "Dispatch Note Date": self.dispatch_date.get().strip(),
                "Notes": self.notes.get().strip()
            }
        filename = self.DISPATCH_INFO_FILE
        # Load existing data first
        all_data = load_from_json(filename)
        
        # Ensure it's a list
        if not isinstance(all_data, list):
            all_data = []

        # Append the new entry
        all_data.append(new_entry)

        # Save everything back
        try:
            with open(filename, "w", encoding="utf-8") as f:
                json.dump(all_data, f, indent=4, ensure_ascii=False)
            messagebox.showinfo("Success", f"Saved dispatch note for {new_entry.get('Customer','Unknown')}")
            print(f"[DEBUG] Saved new dispatch note for customer '{new_entry.get('Customer', '')}'")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save dispatch note:\n{e}")
            print(f"[DEBUG] Error saving JSON: {e}")

        self.append_dispatch_note_excel(new_entry)

    def load_all_dispatch_notes(self):
        """
        Load all dispatch notes safely.
        Ensures return is always a list.
        """
        filename = self.DISPATCH_INFO_FILE
        data = load_from_json(filename)

        if not data:
            print(f"[DEBUG] No dispatch notes found in {filename}. Returning empty list.")
            return []

        if not isinstance(data, list):
            print(f"[DEBUG] Warning: Unexpected data format in {filename}. Resetting to [].")
            return []

        print(f"[DEBUG] Loaded {len(data)} dispatch notes from {filename}.")
        if len(data) > 0:
            sample = data[0].get("Customer", "Unknown")
            print(f"[DEBUG] Example entry - Customer: {sample}")
        return data

    def get_unique_customers(self):
        """
        Return a list of unique customers from saved dispatch notes.
        """
        notes = self.load_all_dispatch_notes()
        customers = list({note.get('Customer', '') for note in notes})
        print(f"[DEBUG] Unique customers: {customers}")
        return customers

    def get_notes_by_customer(self, customer):
        """
        Get all dispatch notes for a specific customer.
        """
        notes = self.load_all_dispatch_notes()
        filtered = [note for note in notes if note.get('Customer', '') == customer]
        print(f"[DEBUG] Found {len(filtered)} notes for customer '{customer}'")
        return filtered

    def load_dispatch_note_to_gui(self, note):
        self.customer_entry.delete(0, tk.END)
        self.customer_entry.insert(0, note.get("Customer", ""))

        self.address_entry.delete(0, tk.END)
        self.address_entry.insert(0, note.get("Address", ""))

        self.phone_entry.delete(0, tk.END)
        self.phone_entry.insert(0, note.get("Phone", ""))

        self.fax_entry.delete(0, tk.END)
        self.fax_entry.insert(0, note.get("Fax", ""))

        self.incharge_entry.delete(0, tk.END)
        self.incharge_entry.insert(0, note.get("Incharge", ""))

        self.contact_number_entry.delete(0, tk.END)
        self.contact_number_entry.insert(0, note.get("Contact Number", ""))

        self.po_ref_entry.delete(0, tk.END)
        self.po_ref_entry.insert(0, note.get("Customer PO Ref", ""))

        self.quotation_entry.delete(0, tk.END)
        self.quotation_entry.insert(0, note.get("Quotation", ""))

        self.project_entry.delete(0, tk.END)
        self.project_entry.insert(0, note.get("Project", ""))

        self.dispatch_date_var.set(note.get("Dispatch Note Date", datetime.now().strftime("%d-%m-%y")))
        self.notes.set(note.get("Notes", ""))

    def on_client_selected(self, event):
        selected_client = self.client_var.get()
        notes = self.get_notes_by_customer(selected_client)
        if notes:
            # Load the first note for this client into the GUI
            self.load_dispatch_note_to_gui(notes[0])

    def append_dispatch_note_excel(self, new_entry):
        """
        Append a single dispatch note to an Excel backup without overwriting previous entries.
        """
        backup_folder = os.path.join(os.getcwd(), "backup")
        os.makedirs(backup_folder, exist_ok=True)
        backup_path = os.path.join(backup_folder, self.EXCEL_BACKUP_FILE)

        df_new = pd.DataFrame([new_entry])

        try:
            # If the backup file exists, append without headers
            if os.path.exists(backup_path):
                df_new.to_excel(backup_path, index=False, header=False, mode='a', engine='openpyxl')
            else:
                # Create a new Excel file with headers
                df_new.to_excel(backup_path, index=False, engine='openpyxl')
            print(f"[DEBUG] Appended dispatch note to Excel backup: {backup_path}")
        except Exception as e:
            print(f"[DEBUG] Failed to append Excel backup: {e}")


    def get_treeview_columns(self):
        """Return columns specific to dispatch notes"""
        return ("Part Number", "Description", "Qty", "Unit Price", "Weight")
    
    def format_item_for_tree(self, product):
        """Format product data for dispatch note tree display"""
        qty = product.get('Qty', 1)
        if isinstance(qty, str) and 'pcs' in qty:
            qty = int(qty.split()[0])
        
        unit_price = product.get('Unit Price', 0)
        total_value = qty * unit_price
        
        return (
            product['Part Number'],
            product['Description'],
            format_qty(qty),
            format_price(unit_price, self.currency_unit.get()),
            #format_price(total_value, self.currency_unit.get()),
            format_weight(product.get('Weight', 0)),
            "Ready",  # Default status for dispatch
            ""  # Empty notes field
        )
    
    def get_export_data(self):
        """Return data formatted for dispatch note export"""
        data = []
        
        # Validate dispatch information
        customer = self.customer_entry.get().strip()
        print(f"[DEBUG] Customer entry: '{customer}'")
        if not customer:
            raise ValueError("Customer name is required")

        dispatch_date = self.dispatch_date.get().strip()
        print(f"[DEBUG] Dispatch date entry: '{dispatch_date}'")
        if not dispatch_date:
            raise ValueError("Dispatch date is required")

        # Debug: check treeview rows
        children = self.item_tree.get_children()
        print(f"[DEBUG] Number of rows in treeview: {len(children)}")
        if not children:
            print("[DEBUG] No items found in treeview!")
            return []

        # Process each item in the tree
        for idx, row in enumerate(children):
            vals = self.item_tree.item(row)['values']
            print(f"[DEBUG] Row {idx} values: {vals} (types: {[type(v) for v in vals]})")

            # Basic validation
            if not vals[0] or not vals[1]:
                print(f"[DEBUG] Skipping row {idx}: missing Part Number or Description")
                continue

            try:
                qty = int(str(vals[2]).split()[0])  # still "1 pcs"
                price = parse_float_from_string(vals[4])  # "USD 108.09" -> 108.09
                weight = parse_float_from_string(vals[5])  # "20.000 kg" -> 20.0
            except Exception as e:
                print(f"[DEBUG] Skipping row {idx} due to parse error: {e}")
                continue


            total_price = qty * price
            total_weight = qty * weight
            
            '''row_data = {
                'Dispatch Date': dispatch_date,
                'Dispatch To': dispatch_to,
                'Destination': self.destination.get(),
                'Transport Method': self.transport_method.get(),
                'Priority': self.priority.get(),
                'Tracking Number': self.tracking_number.get(),
                'Dispatched By': self.dispatched_by.get(),
                'Overall Status': self.dispatch_status.get(),
                'Part Number': vals[0],
                'Description': vals[1],
                'Qty': qty,
                'Unit Price (SAR)': unit_price,
                'Total Value (SAR)': total_value,
                'Unit Weight (kg)': round(weight, 3),
                'Total Weight (kg)': round(total_weight, 3),
                'Item Status': vals[6] if len(vals) > 6 else 'Ready',
                'Item Notes': vals[7] if len(vals) > 7 else '',
                'Special Instructions': self.special_instructions.get()
            }'''

            row_data = {
                'Part Number': vals[0],
                'Description': vals[1],
                'Qty': qty,
                'Unit Price (SAR)': price,
                'Unit Weight (kg)': round(weight, 3),
                'Total Weight (kg)': round(total_weight, 3),
                'Item Status': vals[6] if len(vals) > 6 else 'Ready',
            }
            data.append(row_data)
        
        print(f"[DEBUG] Total exportable rows: {len(data)}")
        return data
    
    def export_template(self):
        """
        Export dispatch note template with required columns (uses default template).
        """
        try:
            export_data = self.get_export_data()

            if not export_data:
                raise ValueError("No data to export. Please add items to the dispatch note.")

            # ✅ Use default template path instead of asking
            default_template = os.path.join(os.getcwd(), "assets", "templates", "dispatch_note_template.docx")

            if not os.path.exists(default_template):
                raise FileNotFoundError(f"Default template not found:\n{default_template}")

            doc = Document(default_template)

            # Mapping of placeholders to actual values
            placeholders = {
                "Customer": self.customer_entry.get().strip(),
                "Address": self.address_entry.get().strip(),
                "Phone_Num": self.phone_entry.get().strip(),
                "Fax": self.fax_entry.get().strip(),
                "Incharge": self.incharge_entry.get().strip(),
                "Customer_PO": self.po_ref_entry.get().strip(),
                "Quotation": self.quotation_entry.get().strip(),
                "Project_Name": self.project_entry.get().strip(),
                "Contact_Num": self.contact_number_entry.get().strip(),
                "Date": self.dispatch_date.get().strip()
            }

            # Replace placeholders in document paragraphs, headers, footers
            for paragraph in doc.paragraphs:
                for key, value in placeholders.items():
                    replace_placeholder_in_paragraph(paragraph, f"{{{key}}}", value)
            for section in doc.sections:
                for header_paragraph in section.header.paragraphs:
                    for key, value in placeholders.items():
                        replace_placeholder_in_paragraph(header_paragraph, f"{{{key}}}", value)
                for footer_paragraph in section.footer.paragraphs:
                    for key, value in placeholders.items():
                        replace_placeholder_in_paragraph(footer_paragraph, f"{{{key}}}", value)

            # Populate item table
            self.populate_item_table(doc, export_data)

            # ✅ Auto-generate filename
            current_date = datetime.now().strftime("%d-%m-%y")
            client_name = self.customer_entry.get().strip() or "Client"
            default_filename = f"DN0{current_date}-{client_name}.docx"

            # ✅ Save automatically into "exports" folder (no dialog)
            export_folder = os.path.join(os.getcwd(), "exports")
            os.makedirs(export_folder, exist_ok=True)
            save_path = os.path.join(export_folder, default_filename)

            doc.save(save_path)
            save_to_json(export_data)
            messagebox.showinfo("Success", f"Dispatch note exported successfully:\n{save_path}")

        except Exception as e:
            messagebox.showerror("Export Failed", f"Failed to export dispatch note:\n{str(e)}")

    def populate_item_table(self, doc, export_data):
        """
        Populate the item table in the Word document.
        
        Args:
            doc: Word document object
            export_data: List of item data dictionaries
        """
        if not doc.tables:
            return

        # Find the correct table (assuming first table with proper headers)
        item_table = None
        for table in doc.tables:
            if len(table.rows) > 0:
                headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
                if 'no.' in headers and 'item' in headers and 'description' in headers:
                    item_table = table
                    break

        if not item_table:
            raise ValueError("Could not find the item table in the Word template.")

        # Clear existing rows (except header)
        for row in item_table.rows[1:]:
            table_element = item_table._tbl
            table_element.remove(row._tr)

        # Add new rows for each item
        for i, item in enumerate(export_data, start=1):
            row_cells = item_table.add_row().cells
            if len(row_cells) >= 4:
                row_cells[0].text = str(i)
                row_cells[1].text = str(item['Part Number'])
                row_cells[2].text = str(item['Description'])
                row_cells[3].text = str(item['Qty'])

    def get_column_width(self, col):
        """Get column width for dispatch note treeview"""
        width_map = {
            "Part Number": 100,
            "Description": 180,
            "Qty": 60,
            "Unit Price": 80,
            "Weight": 80,
        }
        return width_map.get(col, 100)
    
    def print_dispatch_note_pdf(self):
        """
        Export dispatch note as PDF and send to the system printer.
        """
        try:
            # Step 2a: Export data
            export_data = self.get_export_data()
            if not export_data:
                messagebox.showwarning("No Data", "No items to export.")
                return

            # Step 2b: Export as PDF (create temporary PDF file)
            client_name = self.customer_entry.get().strip() or "Client"
            current_date = datetime.now().strftime("%d-%m-%y")
            pdf_filename = f"DN0{current_date}-{client_name}.pdf"
            temp_dir = tempfile.gettempdir()
            pdf_path = os.path.join(temp_dir, pdf_filename)

            # Ask user for Word template
            # ✅ Use default template path instead of asking
            default_template = os.path.join(os.getcwd(), "assets", "templates", "dispatch_note_template.docx")

            if not os.path.exists(default_template):
                raise FileNotFoundError(f"Default template not found:\n{default_template}")

            template_path = default_template
            if not template_path:
                return

            # Fill template with placeholders
            doc = Document(template_path)
            placeholders = {
                "Customer": self.customer_entry.get().strip(),
                "Address": self.address_entry.get().strip(),
                "Phone_Num": self.phone_entry.get().strip(),
                "Fax": self.fax_entry.get().strip(),
                "Incharge": self.incharge_entry.get().strip(),
                "Customer_PO": self.po_ref_entry.get().strip(),
                "Quotation": self.quotation_entry.get().strip(),
                "Project_Name": self.project_entry.get().strip(),
                "Contact_Num": self.contact_number_entry.get().strip(),
                "Date": self.dispatch_date.get().strip()
            }
            for paragraph in doc.paragraphs:
                for key, value in placeholders.items():
                    replace_placeholder_in_paragraph(paragraph, f"{{{key}}}", value)

            self.populate_item_table(doc, export_data)

            # Save temp Word file
            temp_docx = os.path.join(temp_dir, "temp_dispatch_note.docx")
            doc.save(temp_docx)

            # Convert Word to PDF
            convert(temp_docx, pdf_path)

            # Step 2c: Print the PDF using system default viewer
            if os.name == "nt":  # Windows
                os.startfile(pdf_path, "print")
            elif sys.platform == "darwin":  # macOS
                os.system(f"lp '{pdf_path}'")
            else:  # Linux
                os.system(f"lp '{pdf_path}'")

            messagebox.showinfo("Success", f"Dispatch note sent to printer.")

        except Exception as e:
            messagebox.showerror("Print Failed", f"Failed to print dispatch note:\n{e}")



    def on_double_click(self, event):
        """Override to allow editing Notes and Status columns"""
        item = self.item_tree.identify_row(event.y)
        column = self.item_tree.identify_column(event.x)

        if not item:
            return
            
        column_id = int(column[1]) - 1
        column_name = self.item_tree["columns"][column_id]
        
        # Allow editing Qty, Status, and Notes columns
        if column_name not in ["Qty", "Status", "Notes"]:
            return
            
        current_value = self.item_tree.item(item)['values'][column_id]
        
        # Create entry widget for editing
        if column_name == "Status":
            # Use combobox for status
            widget = ttk.Combobox(
                self.item_tree,
                values=["Ready", "Packed", "Shipped", "Delivered", "Issue"],
                state="readonly"
            )
            widget.set(current_value)
        else:
            # Regular entry for other fields
            widget = ttk.Entry(self.item_tree)
            widget.insert(0, current_value)
        
        bbox = self.item_tree.bbox(item, column)
        widget.place(x=bbox[0], y=bbox[1], width=bbox[2], height=bbox[3])
        
        def on_enter(e):
            new_value = widget.get().strip()
            try:
                if column_name == "Qty":
                    qty_str = ''.join(filter(str.isdigit, new_value))
                    if not qty_str:
                        tk.messagebox.showerror("Invalid Input", "Quantity must be a positive number")
                        return
                    
                    qty_val = int(qty_str)
                    if qty_val <= 0:
                        tk.messagebox.showerror("Invalid Input", "Quantity must be greater than 0")
                        return
                    
                    new_value = format_qty(qty_val)
                
                values = list(self.item_tree.item(item)['values'])
                values[column_id] = new_value
                self.item_tree.item(item, values=values)
                
            except ValueError:
                tk.messagebox.showerror("Invalid Input", f"Please enter a valid {column_name.lower()}")
            finally:
                widget.destroy()
        
        widget.bind('<Return>', on_enter)
        widget.bind('<Escape>', lambda e: widget.destroy())
        widget.focus_set()


# Test the module independently
if __name__ == "__main__":
    # Create a simple test window
    root = tk.Tk()
    root.withdraw()  # Hide root window
    
    try:
        app = DispatchNoteGenerator(root)
        app.mainloop()
    except Exception as e:
        print(f"Error running dispatch note generator: {e}")
        messagebox.showerror("Error", f"Failed to start application:\n{e}")
    finally:
        root.destroy()