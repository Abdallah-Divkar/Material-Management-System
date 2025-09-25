"""
Material List Generator Module

This module provides functionality for generating material lists with proper
title section layout and consistent styling throughout the application.

Author: Material Management System
Version:  1.0
Last Modified: 2025-09-17
"""

import tkinter as tk
from tkinter import ttk, messagebox
from datetime import datetime
import sys
import os
import tempfile
import json
import tkinter.filedialog as fd
from datetime import datetime
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx2pdf import convert
import pandas as pd
from common.utils import replace_placeholder_in_paragraph, save_to_json, load_from_json, parse_float_from_string

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from common.base_generator import BaseGenerator
from common.utils import format_qty, format_price, format_weight, format_currency
from common.currency_handler import CurrencyHandler


class MaterialListGenerator(BaseGenerator):
    """
    Material List Generator - for incoming material deliveries

    This class extends BaseGenerator to provide material list specific functionality
    with improved title section layout and consistent UI structure.
    """
    def __init__(self, parent):
        """Initialize the Material List Generator with proper title configuration."""
        self.module_title = "Material List Generator"
        self.export_button_text = "Export Material List"
        self.MATERIAL_INFO_FILE = "./backup/client_info_cache.json"
        self.EXCEL_BACKUP_FILE = "./backup/client_info_backup.xlsx"
        #self.default_filename = f"material_list_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"

        # Initialize BaseGenerator and self.main_frame
        super().__init__(parent, self.module_title)

        self.delivery_date_var = tk.StringVar(value=datetime.now().strftime("%d-%m-%y"))
        self.notes = tk.StringVar()

        # Create the title section with material info callback
        self.create_title_section(left_frame_callback=self.create_material_info_inline)
        #self.load_material_info()


        # Create custom widgets
        self.create_custom_widgets()

    def create_title_section(self, left_frame_callback=None):
        """
        Create an improved title section with proper layout structure.
        
        Args:
            left_frame_callback: Optional callback for creating left frame content
        """
        # Create main header frame with consistent styling
        header_frame = tk.Frame(self.main_frame, bg="#00A695", height=120)
        header_frame.grid(row=0, column=0, sticky="ew", padx=15, pady=(5, 0))
        
        # Configure column weights for responsive layout
        header_frame.grid_columnconfigure(0, weight=1)  # Left side (material info)
        header_frame.grid_columnconfigure(1, weight=0)  # Right side (logo + title)
        
        # Store header frame reference for child classes
        self.header_frame = header_frame
        
        # Left side: Material Information (if callback provided)
        if left_frame_callback:
            self.material_info_frame = left_frame_callback(header_frame)
        else:
            # Default empty frame if no callback provided
            self.material_info_frame = tk.Frame(header_frame, bg="#F0F0F0")
            self.material_info_frame.grid(row=0, column=0, sticky="nsew", padx=(0, 10))

        # Right side: Logo and Title section
        self.create_logo_title_section(header_frame)

    def create_logo_title_section(self, parent_frame):
        """
        Create the logo and title section with consistent styling.
        
        Args:
            parent_frame: Parent frame to contain logo and title elements
        """
        logo_title_frame = tk.Frame(parent_frame, bg="#00A695")
        logo_title_frame.grid(row=0, column=1, sticky="ne", padx=(10, 0), pady=10)
        
        # Logo section
        try:
            logo_path = os.path.join("assets", "images" "mts_logo.png")
            if os.path.exists(logo_path):
                from PIL import Image, ImageTk
                logo_image = Image.open(logo_path).resize((80, 80), Image.Resampling.LANCZOS)
                self.logo_photo = ImageTk.PhotoImage(logo_image)
                logo_label = tk.Label(
                    logo_title_frame, 
                    image=self.logo_photo, 
                    bg="#00A695",
                    relief="flat"
                )
                logo_label.pack(side="left", padx=(0, 15), pady=5)
        except Exception as e:
            print(f"Warning: Could not load logo - {e}")
            # Create placeholder for logo space
            placeholder = tk.Frame(logo_title_frame, width=80, height=80, bg="#00A695")
            placeholder.pack(side="left", padx=(0, 15), pady=5)
        
        # Title section
        title_text = getattr(self, 'module_title', 'Generator')
        title_label = tk.Label(
            logo_title_frame,
            text=title_text,
            bg="#00A695",
            fg="white",
            font=("Arial", 24, "bold"),
            anchor="e",
            justify="right"
        )
        title_label.pack(side="left", pady=5)

    def create_material_info_inline(self, parent_frame):
        """
        Create compact material information section for the header
        with updated field names and empty entries.
        
        Args:
            parent_frame: Parent frame to contain material info

        Returns:
            tk.LabelFrame: The created material info frame
        """
        info_frame = tk.LabelFrame(
            parent_frame,
            text="Material Information",
            bg="#F0F0F0",
            font=("Arial", 11, "bold"),
            labelanchor="nw",
            relief="raised",
            bd=2
        )
        info_frame.grid(row=0, column=0, sticky="ew")
        
        # Configure grid weights for responsive layout
        for i in range(4):
            info_frame.grid_columnconfigure(i, weight=1)

        self.create_info_field(info_frame, "Select Project:", 0, 0)
        self.client_var = tk.StringVar()
        self.client_dropdown = ttk.Combobox(info_frame, textvariable=self.client_var, width=20)
        self.client_dropdown.grid(row=0, column=1, padx=5, pady=2)

        # Populate with unique clients
        self.client_dropdown['values'] = self.get_unique_customers()

        # Bind selection to load first dispatch note for this client
        self.client_dropdown.bind("<<ComboboxSelected>>", self.on_client_selected)
        
        self.create_info_field(info_frame, "Project:", 1, 0)
        self.project_entry = self.create_info_entry(info_frame, "", 1, 1)
        self.create_info_field(info_frame, "Work Order No:", 1, 2)
        self.po_ref_entry = self.create_info_entry(info_frame, "", 1, 3)
        self.create_info_field(info_frame, "Client:", 2, 0)
        self.incharge_entry = self.create_info_entry(info_frame, "", 2, 1)
        self.create_info_field(info_frame, "Delivery Date:", 2, 2)
        self.delivery_date = self.create_info_entry(info_frame, self.delivery_date_var.get(), 2, 3, textvariable=self.delivery_date_var)

        return info_frame

    def create_info_field(self, parent, text, row, col, width=12):
        """
        Create a consistent label field for material information.
        
        Args:
            parent: Parent widget
            text: Label text
            row: Grid row
            col: Grid column
            width: Label width
        """
        label = tk.Label(
            parent,
            text=text,
            bg="#F0F0F0",
            font=("Arial", 9, "normal"),
            anchor="w",
            width=width
        )
        label.grid(row=row, column=col, sticky="w", padx=1, pady=1)
        return label

    def create_info_entry(self, parent, default_value="", row=0, col=0, textvariable=None, width=15):
        """
        Create a consistent entry field for material information.
        
        Args:
            parent: Parent widget
            default_value: Default entry value
            row: Grid row
            col: Grid column
            textvariable: Optional StringVar for the entry
            width: Entry width
            
        Returns:
            tk.Entry: The created entry widget
        """
        entry = tk.Entry(
            parent,
            font=("Arial", 9),
            width=width,
            relief="sunken",
            bd=1,
            textvariable=textvariable
        )
        if not textvariable and default_value:
            entry.insert(0, default_value)
        entry.grid(row=row, column=col, sticky="ew", padx=1, pady=1)
        return entry

    def _on_mousewheel(self, event):
        self.canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
    
    def create_custom_widgets(self):
        self.get_treeview_columns()
        save_btn = tk.Button(self.btn_frame, text="Save Material Release Form", command=self.save_material_list)
        save_btn.grid(row=0, column=5, sticky="w", padx=10, pady=5)

        load_btn = tk.Button(self.btn_frame, text="Load Material Release Form", command=self.load_all_material_lists)
        load_btn.grid(row=0, column=6, sticky="w", padx=10, pady=5)

    def save_material_list(self, new_entry=None):
        """
        Save a new material list to JSON without overwriting existing notes.
        """
        if new_entry is None:
            new_entry = {
                "Incharge": self.incharge_entry.get().strip(),
                "Customer PO Ref": self.po_ref_entry.get().strip(),
                "Project": self.project_entry.get().strip(),
                "Delivery Note Date": self.delivery_date.get().strip(),
                "Notes": self.notes.get().strip()
            }
        filename = self.MATERIAL_INFO_FILE
        # Load existing data first
        all_data = load_from_json(filename)
        
        # Ensure it's a list
        if not isinstance(all_data, list):
            all_data = []

        # Append the new entry
        all_data.append(new_entry)

        # Save everything back
        try:
            with open(filename, "w", encoding="utf-8") as f:
                json.dump(all_data, f, indent=4, ensure_ascii=False)
            messagebox.showinfo("Success", f"Saved material list for {new_entry.get('Customer','Unknown')}")
            print(f"[DEBUG] Saved new material list for customer '{new_entry.get('Customer', '')}'")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save material list:\n{e}")
            print(f"[DEBUG] Error saving JSON: {e}")

        self.append_material_list_excel(new_entry)

    def load_all_material_lists(self):
        """
        Load all material lists safely.
        Ensures return is always a list.
        """
        filename = self.MATERIAL_INFO_FILE
        data = load_from_json(filename)

        if not data:
            print(f"[DEBUG] No material lists found in {filename}. Returning empty list.")
            return []

        if not isinstance(data, list):
            print(f"[DEBUG] Warning: Unexpected data format in {filename}. Resetting to [].")
            return []

        print(f"[DEBUG] Loaded {len(data)} material lists from {filename}.")
        if len(data) > 0:
            sample = data[0].get("Customer", "Unknown")
            print(f"[DEBUG] Example entry - Customer: {sample}")
        return data

    def get_unique_customers(self):
        """
        Return a list of unique customers from saved material lists.
        """
        notes = self.load_all_material_lists()
        customers = list({note.get('Customer', '') for note in notes})
        print(f"[DEBUG] Unique customers: {customers}")
        return customers

    def get_notes_by_customer(self, customer):
        """
        Get all material lists for a specific customer.
        """
        notes = self.load_all_material_lists()
        filtered = [note for note in notes if note.get('Customer', '') == customer]
        print(f"[DEBUG] Found {len(filtered)} notes for customer '{customer}'")
        return filtered

    def load_material_list_to_gui(self, note):
        '''
        self.address_entry.delete(0, tk.END)
        self.address_entry.insert(0, note.get("Address", ""))

        self.phone_entry.delete(0, tk.END)
        self.phone_entry.insert(0, note.get("Phone", ""))

        self.fax_entry.delete(0, tk.END)
        self.fax_entry.insert(0, note.get("Fax", ""))'''

        self.incharge_entry.delete(0, tk.END)
        self.incharge_entry.insert(0, note.get("Incharge", ""))

        '''self.contact_number_entry.delete(0, tk.END)
        self.contact_number_entry.insert(0, note.get("Contact Number", ""))'''

        self.po_ref_entry.delete(0, tk.END)
        self.po_ref_entry.insert(0, note.get("Customer PO Ref", ""))

        '''self.quotation_entry.delete(0, tk.END)
        self.quotation_entry.insert(0, note.get("Quotation", ""))'''

        self.project_entry.delete(0, tk.END)
        self.project_entry.insert(0, note.get("Project", ""))

        self.delivery_date_var.set(note.get("Delivery Note Date", datetime.now().strftime("%d-%m-%y")))
        self.notes.set(note.get("Notes", ""))

    def on_client_selected(self, event):
        selected_client = self.client_var.get()
        notes = self.get_notes_by_customer(selected_client)
        if notes:
            # Load the first note for this client into the GUI
            self.load_material_list_to_gui(notes[0])

    def append_material_list_excel(self, new_entry):
        """
        Append a single material list to an Excel backup without overwriting previous entries.
        """
        backup_folder = os.path.join(os.getcwd(), "backup")
        os.makedirs(backup_folder, exist_ok=True)
        backup_path = os.path.join(backup_folder, self.EXCEL_BACKUP_FILE)

        df_new = pd.DataFrame([new_entry])

        try:
            # If the backup file exists, append without headers
            if os.path.exists(backup_path):
                df_new.to_excel(backup_path, index=False, header=False, mode='a', engine='openpyxl')
            else:
                # Create a new Excel file with headers
                df_new.to_excel(backup_path, index=False, engine='openpyxl')
            print(f"[DEBUG] Appended material list to Excel backup: {backup_path}")
        except Exception as e:
            print(f"[DEBUG] Failed to append Excel backup: {e}")


    def get_treeview_columns(self):
        """Return columns specific to material lists"""
        return ("Part Number", "Description", "Supplier", "Qty", "Unit Price", "Total Price", "Weight")
    
    def format_item_for_tree(self, product):
        """Format product data for material list tree display"""
        qty = int(str(product.get('Qty', 1)).split()[0]) if 'pcs' in str(product.get('Qty', 1)) else product.get('Qty', 1)
        unit_price = product.get('Unit Price', 0)
        total_price = qty * unit_price
        weight = product.get('Weight', 0)
        total_weight = qty * weight

        return (
            product.get('Part Number', ''),
            product.get('Description', ''),
            product.get('Supplier', ''),
            format_qty(qty),
            format_currency(unit_price, self.currency_unit.get()),   # ✅ with currency
            format_currency(total_price, self.currency_unit.get()),
            format_weight(weight),
        )

    def get_export_data(self):
        """Return data formatted for export with column mapping by name"""
        export_data = []

        # Get current columns from treeview
        tree_columns = self.item_tree["columns"]

        for row_id in self.item_tree.get_children():
            row = self.item_tree.item(row_id, "values")

            # Build dict mapping col_name -> value
            row_dict = {col: row[idx] if idx < len(row) else "" for idx, col in enumerate(tree_columns)}

            export_data.append(row_dict)

        return export_data

    def export_template(self):
        """
        Export material list template with required columns (uses default template).
        """
        try:
            export_data = self.get_export_data()

            if not export_data:
                raise ValueError("No data to export. Please add items to the material list.")

            # ✅ Use default template path instead of asking
            default_template = os.path.join(os.getcwd(), "assets", "templates", "mrf_template.docx")

            if not os.path.exists(default_template):
                raise FileNotFoundError(f"Default template not found:\n{default_template}")

            doc = Document(default_template)

            # Mapping of placeholders to actual values
            placeholders = {
                "Incharge": self.incharge_entry.get().strip(),
                "Customer_PO": self.po_ref_entry.get().strip(),
                "Date": self.delivery_date.get().strip(),
                "Project_Name": self.project_entry.get().strip(),
            }

            # Replace placeholders in document paragraphs, headers, footers
            for paragraph in doc.paragraphs:
                for key, value in placeholders.items():
                    replace_placeholder_in_paragraph(paragraph, f"{{{key}}}", value)
            for section in doc.sections:
                for header_paragraph in section.header.paragraphs:
                    for key, value in placeholders.items():
                        replace_placeholder_in_paragraph(header_paragraph, f"{{{key}}}", value)
                for footer_paragraph in section.footer.paragraphs:
                    for key, value in placeholders.items():
                        replace_placeholder_in_paragraph(footer_paragraph, f"{{{key}}}", value)

            # Populate item table
            self.populate_item_table(doc, export_data)

            # ✅ Auto-generate filename
            current_date = datetime.now().strftime("%d-%m-%y")
            incharge_name = self.incharge_entry.get().strip() or "Incharge"
            project_name = self.project_entry.get().strip() or "Project"
            default_filename = f"MRF-{current_date}-{project_name}-{incharge_name}.docx"

            # ✅ Save automatically into "exports" folder (no dialog)
            export_folder = os.path.join(os.getcwd(), "exports")
            os.makedirs(export_folder, exist_ok=True)
            save_path = os.path.join(export_folder, default_filename)

            doc.save(save_path)
            save_to_json(export_data)
            messagebox.showinfo("Success", f"Material list exported successfully:\n{save_path}")

        except Exception as e:
            messagebox.showerror("Export Failed", f"Failed to export material list:\n{str(e)}")

    def populate_item_table(self, doc, export_data):
        """
        Populate the item table in the Word document.
        
        Args:
            doc: Word document object
            export_data: List of item data dictionaries
        """
        if not doc.tables:
            return

        # Find the correct table (assuming first table with proper headers)
        item_table = None
        for table in doc.tables:
            if len(table.rows) > 0:
                headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
                if 'item' in headers and 'part number' in headers and 'description' in headers and 'qty' in headers:
                    item_table = table
                    break

        if not item_table:
            raise ValueError("Could not find the item table in the Word template.")

        # Clear existing rows (except header)
        for row in item_table.rows[1:]:
            table_element = item_table._tbl
            table_element.remove(row._tr)

        # Add new rows for each item
        for i, item in enumerate(export_data, start=1):
            row_cells = item_table.add_row().cells
            if len(row_cells) >= 4:
                row_cells[0].text = str(i)
                row_cells[1].text = str(item['Part Number'])
                row_cells[2].text = str(item['Description'])
                row_cells[3].text = str(item['Qty'])

    def get_column_width(self, col):
        """
        Get column width for material list treeview.
        
        Args:
            col: Column name
            
        Returns:
            int: Column width in pixels
        """
        width_map = {
            "Item": 2,
            "Part Number": 50,
            "Description": 200,
            "Qty": 3,
        }
        return width_map.get(col, 100)
    
    def print_material_list_pdf(self):
        """
        Export material list as PDF and send to the system printer.
        """
        try:
            # Step 2a: Export data
            export_data = self.get_export_data()
            if not export_data:
                messagebox.showwarning("No Data", "No items to export.")
                return

            # Step 2b: Export as PDF (create temporary PDF file)
            incharge_name = self.incharge_entry.get().strip() or "Incharge"
            project_name = self.project_entry.get().strip() or "Project"
            current_date = datetime.now().strftime("%d-%m-%y")
            pdf_filename = f"MRF-{current_date}-{project_name}-{incharge_name}.pdf"
            temp_dir = tempfile.gettempdir()
            pdf_path = os.path.join(temp_dir, pdf_filename)

            # Ask user for Word template
            from tkinter import filedialog as fd

            # ✅ Use default template path instead of asking
            default_template = os.path.join(os.getcwd(), "assets", "templates", "mrf_template.docx")

            if not os.path.exists(default_template):
                raise FileNotFoundError(f"Default template not found:\n{default_template}")

            doc = Document(default_template)

            template_path = default_template
            if not template_path:
                return

            # Fill template with placeholders
            doc = Document(template_path)
            placeholders = {
                "Incharge": self.incharge_entry.get().strip(),
                "Customer_PO": self.po_ref_entry.get().strip(),
                "Project_Name": self.project_entry.get().strip(),
                "Date": self.delivery_date.get().strip()
            }
            for paragraph in doc.paragraphs:
                for key, value in placeholders.items():
                    replace_placeholder_in_paragraph(paragraph, f"{{{key}}}", value)

            self.populate_item_table(doc, export_data)

            # Save temp Word file
            temp_docx = os.path.join(temp_dir, "temp_material_list.docx")
            doc.save(temp_docx)

            # Convert Word to PDF
            convert(temp_docx, pdf_path)

            # Step 2c: Print the PDF using system default viewer
            if os.name == "nt":  # Windows
                os.startfile(pdf_path, "print")
            elif sys.platform == "darwin":  # macOS
                os.system(f"lp '{pdf_path}'")
            else:  # Linux
                os.system(f"lp '{pdf_path}'")

            messagebox.showinfo("Success", f"Material list sent to printer.")

        except Exception as e:
            messagebox.showerror("Print Failed", f"Failed to print material list:\n{e}")


# Test the module independently
if __name__ == "__main__":
    # Create a simple test window
    root = tk.Tk()
    root.withdraw()  # Hide root window
    
    try:
        app = MaterialListGenerator(root)
        app.mainloop()
    except Exception as e:
        print(f"Error running material list generator: {e}")
        messagebox.showerror("Error", f"Failed to start application:\n{e}")
    finally:
        root.destroy()