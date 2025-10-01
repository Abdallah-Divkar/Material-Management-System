# common/base_generator.py
import json
import os
from docx import Document
import requests
from dotenv import load_dotenv
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from PIL import Image, ImageTk
import pandas as pd
from abc import ABC, abstractmethod
import pandas as pd
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter

from common.excel_handler import get_products, get_product_details, save_products_cache
from common.utils import format_qty, format_price, format_weight, format_currency, load_from_json, parse_float_from_string
from common.currency_handler import CurrencyHandler

load_dotenv()

API_KEY = os.getenv("EXCHANGE_RATE_API_KEY")


class BaseGenerator(tk.Toplevel, ABC):
    """Base class for all generator modules with improved layout structure"""
    
    def __init__(self, parent, title="Generator"):
        super().__init__(parent)
        self.parent = parent
        self.title(title)
        self.geometry("1200x900")  # Increased size for better layout
        self.configure(bg="#00A695")
        
        # Common attributes
        self.products = get_products()
        self.currency_unit = tk.StringVar(value="SAR")
        self.selected_items = []
        self.combo_display_list = []
        
        # Build display list for combobox
        self.build_combo_display_list()
        
        # Create the UI structure
        self.create_base_widgets()
        
        # Focus on this window
        self.focus_set()
        self.grab_set()
    
    def build_combo_display_list(self):
        """Build the display list for the combobox"""
        self.combo_display_list = [
            f"{p['Part Number']} - {p['Description']}" 
            for p in self.products
        ]
    
    def create_base_widgets(self):
        """Create the base UI components common to all generators"""
        
        # Create main canvas and scrollbar for the entire window
        self.main_canvas = tk.Canvas(self, bg="#00A695", highlightthickness=0)
        self.main_scrollbar = ttk.Scrollbar(self, orient="vertical", command=self.main_canvas.yview)
        self.main_canvas.configure(yscrollcommand=self.main_scrollbar.set)

        # Pack the main canvas and scrollbar
        self.main_canvas.pack(side="left", fill="both", expand=True, padx=15, pady=5)
        self.main_scrollbar.pack(side="right", fill="y", pady=15)

        # Create main frame inside the canvas
        self.main_frame = tk.Frame(self.main_canvas, bg="#00A695")
        self.main_canvas_frame = self.main_canvas.create_window((0, 0), window=self.main_frame, anchor="nw")

        # Configure main frame grid
        self.main_frame.grid_columnconfigure(0, weight=1)

        # Bind canvas configurations
        self.main_frame.bind("<Configure>", self.on_main_frame_configure)
        self.main_canvas.bind("<Configure>", self.on_main_canvas_configure)
        
        # NOTE: Title section is now created by subclasses via create_title_section()
        # This allows for customization while maintaining consistent base structure
        
        # Search section
        self.create_search_section()
        
        # Details section
        self.create_details_section()
        
        # Item management section
        #self.create_item_management_section()
        
        # Items treeview
        self.create_items_treeview()
        
        # Control buttons
        self.create_control_buttons()

    def create_search_section(self):
        """Create the search/product selection section"""
        search_frame = tk.Frame(self.main_frame, bg="#00A695")
        search_frame.grid(row=1, column=0, pady=10, sticky="ew", padx=15)

        tk.Label(
            search_frame, 
            text="Search Products (Part Number, Description):",
            bg="#00A695", 
            fg="white", 
            font=("Arial", 12)
        ).pack(side="left", padx=(0, 10))

        self.combo_var = tk.StringVar()
        self.combo = ttk.Combobox(
            search_frame, 
            values=self.combo_display_list, 
            textvariable=self.combo_var, 
            width=50
        )
        self.combo.pack(side="left", pady=5)
        self.combo.bind('<KeyRelease>', self.on_keyrelease)
        self.combo.bind('<<ComboboxSelected>>', self.on_item_selected)
        self.combo.bind('<Return>', self.on_enter_pressed)

        # Upload File button
        ttk.Button(search_frame, text="Upload File", command=self.upload_file).pack(side="left", padx=10)

        # Currency switcher
        self.currency_var = tk.StringVar(value="SAR")
        currency_combo = ttk.Combobox(
            search_frame,
            values=["SAR", "USD"],
            textvariable=self.currency_var,
            width=5,
            state="readonly"
        )
        currency_combo.pack(side="left", padx=(10,0))
        currency_combo.bind("<<ComboboxSelected>>", self.on_currency_changed)

    def create_details_section(self):
        """Create the product details section with improved layout"""
        # Create container frame for both details and treeview
        self.container_frame = tk.Frame(self.main_frame, bg="#F0F0F0")
        self.container_frame.grid(row=2, column=0, sticky="nsew", padx=15, pady=10)
        
        # Configure container frame grid weights
        self.main_frame.grid_rowconfigure(2, weight=1)

        # Left side - Details section
        details_container = tk.Frame(self.container_frame, bg="#F0F0F0")
        details_container.pack(side='left', fill='both', padx=(0, 10))

        # Details frame with fixed width
        outer_frame = tk.Frame(details_container, bg="#F0F0F0", width=400, height=500)
        outer_frame.pack(fill='both', expand=True)
        outer_frame.propagate(False)

        # Add canvas and scrollbar for details
        self.canvas = tk.Canvas(outer_frame, bg="#F0F0F0", width=400)
        scrollbar = ttk.Scrollbar(outer_frame, orient="vertical", command=self.canvas.yview)

        # Detail frame (managed by canvas only)
        self.detail_frame = tk.Frame(self.canvas, bg="#F0F0F0", relief="sunken", borderwidth=2)

        # Configure canvas
        self.canvas.configure(yscrollcommand=scrollbar.set)
        self.canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        # Create window in canvas for detail_frame
        self.canvas_frame = self.canvas.create_window((0, 0), window=self.detail_frame, anchor="nw")

        # Bind canvas configurations
        self.detail_frame.bind("<Configure>", self.on_frame_configure)
        self.canvas.bind("<Configure>", self.on_canvas_configure)

        # Add middle buttons section
        self.create_middle_buttons()

        self.detail_labels = {}

    def create_middle_buttons(self):
        """Create buttons column between details and treeview"""
        buttons_frame = tk.Frame(self.container_frame, bg="#F0F0F0", width=50)
        buttons_frame.pack(side='left', fill='y', padx=10)
        
        # Add spacing at top
        tk.Frame(buttons_frame, height=20, bg="#F0F0F0").pack()
        
        # Add Selected Items button "+"
        self.add_btn = ttk.Button(
            buttons_frame,
            text="+",
            width=3,
            command=self.add_selected_items
        )
        self.add_btn.pack(pady=5)
        
        # Remove Selected Items button "-"
        self.remove_btn = ttk.Button(
            buttons_frame,
            text="-",
            width=3,
            command=self.remove_selected_item,
            state='disabled'  # Disabled initially, enabled when items exist
        )
        self.remove_btn.pack(pady=5)

        # Select All button
        self.select_all_btn = ttk.Button(
            buttons_frame,
            text="All",
            width=3,
            command=self.select_all_items
        )
        self.select_all_btn.pack(pady=5)

        
        # Move Up button
        self.move_up_btn = ttk.Button(
            buttons_frame,
            text="↑",
            width=3,
            command=self.move_item_up
        )
        self.move_up_btn.pack(pady=5)
        
        # Move Down button
        self.move_down_btn = ttk.Button(
            buttons_frame,
            text="↓",
            width=3,
            command=self.move_item_down
        )
        self.move_down_btn.pack(pady=5)

        # Clear Details button
        self.clear_details_btn = ttk.Button(
            buttons_frame, 
            text="Clear Details",
            width=12,
            command=self.clear_details
        )
        self.clear_details_btn.pack(pady=5)

    def move_item_up(self):
        """Move selected item up in the treeview"""
        selected = self.item_tree.selection()
        if not selected:
            return
            
        for item in selected:
            idx = self.item_tree.index(item)
            if idx > 0:
                prev = self.item_tree.prev(item)
                self.item_tree.move(item, self.item_tree.parent(item), self.item_tree.index(prev))

    def move_item_down(self):
        """Move selected item down in the treeview"""
        selected = self.item_tree.selection()
        if not selected:
            return
            
        for item in reversed(selected):
            next_item = self.item_tree.next(item)
            if next_item:
                self.item_tree.move(item, self.item_tree.parent(item), self.item_tree.index(next_item))

    def create_items_treeview(self):
        """Create the items treeview - to be customized by subclasses"""
        columns = self.get_treeview_columns()
        
        # Create a frame to hold the Treeview and Scrollbar
        tree_frame = ttk.Frame(self.container_frame)
        tree_frame.pack(side='left', fill="both", expand=True)

        # Create vertical scrollbar
        tree_scrollbar = ttk.Scrollbar(tree_frame, orient="vertical")
        tree_scrollbar.pack(side="right", fill="y")

        # Create Treeview and attach scrollbar
        self.item_tree = ttk.Treeview(
            tree_frame,
            columns=columns,
            show='headings',
            yscrollcommand=tree_scrollbar.set
        )
        self.item_tree.pack(side="left", fill="both", expand=True)

        # Configure scrollbar
        tree_scrollbar.config(command=self.item_tree.yview)

        # Set up headings and column widths
        for col in columns:
            self.item_tree.heading(col, text=col)
            width = self.get_column_width(col)
            self.item_tree.column(col, width=width)

        # Add double-click binding for editing
        self.item_tree.bind('<Double-1>', self.on_double_click)

    def create_control_buttons(self):
        """Create the control buttons section"""
        btn_frame = tk.Frame(self.main_frame, bg="#00A695")
        btn_frame.grid(row=4, column=0, pady=5, sticky="ew", padx=15)

        ttk.Button(btn_frame, text="Reset", command=self.reset_items).grid(row=0, column=1, padx=10)
        
        # Export button text customizable by subclass
        ttk.Button(btn_frame, text='Export to Excel', command=self.export_to_excel).grid(row=0, column=2, padx=10)
        
        if hasattr(self, 'export_template'):
            export_text = getattr(self, 'export_button_text', 'Export as Template')
            ttk.Button(btn_frame, text=export_text, command=self.export_template).grid(row=0, column=3, padx=10)

        # Add Print button if subclass has print_delivery_note_pdf
        if hasattr(self, 'print_delivery_note_pdf'):
            ttk.Button(btn_frame, text="Print", command=self.print_delivery_note_pdf).grid(row=0, column=4, padx=10)
        elif hasattr(self, 'print_dispatch_note_pdf'):
            ttk.Button(btn_frame, text="Print", command=self.print_dispatch_note_pdf).grid(row=0, column=4, padx=10)
        elif hasattr(self, 'print_material_list_pdf'):
            ttk.Button(btn_frame, text="Print", command=self.print_material_list_pdf).grid(row=0, column=4, padx=10)

        '''save_btn = tk.Button(btn_frame, text="Save Delivery Info", command=self.save_delivery_note)
        save_btn.grid(row=0, column=5, sticky="w", padx=10, pady=5)

        load_btn = tk.Button(btn_frame, text="Load Delivery Info", command=self.load_all_delivery_notes())
        load_btn.grid(row=0, column=6, sticky="w", padx=10, pady=5)'''

        ttk.Button(btn_frame, text="Home", command=self.return_home).grid(row=0, column=7, padx=10)

        self.btn_frame = btn_frame

    '''def update_client_info_cache(self, customer, project, address="", phone="", incharge=""):
        cache_file = "./backup/client_info_cache.json"
        data = load_from_json(cache_file)
        if not isinstance(data, list):
            data = []

        # Avoid duplicate entries
        if not any(d.get("Customer") == customer and d.get("Project") == project for d in data):
            data.append({
                "Customer": customer,
                "Project": project,
                "Address": address,
                "Phone": phone,
                "Incharge": incharge
            })
            with open(cache_file, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=4, ensure_ascii=False)'''
    
    def update_client_info_cache(self,
        delivery_no,
        customer,
        project="",
        address="",
        phone="",
        incharge="",
        contact_number="",
        po_ref="",
        quotation="",
        subject="",
        delivery_date=""):
        """
        Update the client info cache JSON with only template placeholders.
        Avoids duplicates (same customer + project + delivery_no).

        Args:
            delivery_no (str): Delivery / Dispatch / Material release number
            customer (str): Customer name
            project (str): Project name
            address (str): Address
            phone (str): Phone number
            incharge (str): Attn / Incharge
            contact_number (str): Contact number
            po_ref (str): Customer PO reference
            quotation (str): Quotation reference
            subject (str): Subject
            delivery_date (str): Delivery date
        """
        cache_file = "./backup/client_info_cache.json"
        data = load_from_json(cache_file)
        if not isinstance(data, list):
            data = []

        # Avoid duplicate entries: match on customer + project + delivery_no
        if not any(
            d.get("Customer") == customer
            and d.get("Project_Name") == project
            and d.get("Delivery_No") == delivery_no
            for d in data
        ):
            data.append({
                "Customer": customer,
                "Project_Name": project,
                "Address": address,
                "Phone_Num": phone,
                "Incharge": incharge,
                "Contact_Num": contact_number,
                "Customer_PO": po_ref,
                "Quotation": quotation,
                "Subject": subject,
                "Date": delivery_date,
                "Delivery_No": delivery_no  # can be Delivery / Dispatch / Material release number
            })
            with open(cache_file, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=4, ensure_ascii=False)
            print(f"[DEBUG] Updated client info cache for {customer} - {project}")



    # Abstract methods to be implemented by subclasses
    @abstractmethod
    def create_custom_widgets(self):
        """Create custom widgets specific to the generator type"""
        pass
    
    @abstractmethod
    def get_treeview_columns(self):
        """Return the columns for the treeview"""
        pass
    
    @abstractmethod
    def get_export_data(self):
        """Return data formatted for export"""
        pass
    
    # Common functionality methods
    def get_column_width(self, col):
        """Get column width for treeview"""
        width_map = {
            "Part Number": 100,
            "Description": 200,
            "Qty": 80,
            "Supplier": 100,
            "Unit Price": 100,
            "Weight": 80,
            "Customer": 120,
            "Delivery Date": 100,
            "Location": 120,
            "Notes": 150,
            "Status": 80
        }
        return width_map.get(col, 100)
    
    def on_keyrelease(self, event):
        typed = self.combo_var.get().lower()
        
        # Disable add button while typing
        if hasattr(self, 'add_btn'):
            self.add_btn.config(state='disabled')
        
        if not typed:
            # Show full list silently (no popup)
            self.combo['values'] = self.combo_display_list
            for child in self.detail_frame.winfo_children():
                child.destroy()
            self.detail_labels = {}
            return

        # Filter products
        filtered_products = [
            p for p in self.products
            if typed in str(p['Part Number']).lower() or typed in str(p['Description']).lower()
        ]
        
        filtered_display = [
            f"{p['Part Number']} - {p['Description']}" for p in filtered_products
        ]
        self.combo['values'] = filtered_display

        # Clear previous details silently
        for child in self.detail_frame.winfo_children():
            child.destroy()
        self.detail_labels = {}

        # Show filtered products
        for product in filtered_products:
            self.show_details(product)
    
    def on_item_selected(self, event):
        selection = self.combo_var.get()
        if not selection:
            return

        part_number = str(selection.split(' - ')[0].strip())
        product = next((p for p in self.products if str(p.get('Part Number', '')).strip() == part_number), None)
        
        if product:
            self.show_details(product)
            if hasattr(self, 'add_btn'):
                self.add_btn.config(state='normal')
        else:
            messagebox.showerror("Error", f"Product with part number '{part_number}' not found.")
            if hasattr(self, 'add_btn'):
                self.add_btn.config(state='disabled')

    def clear_details(self):
        """Clear all widgets in detail frame"""
        if self.detail_frame.winfo_children():
            if messagebox.askyesno("Confirm Clear", "Are you sure you want to clear all details?"):
                for child in self.detail_frame.winfo_children():
                    child.destroy()
                
                self.detail_labels = {}
                self.current_product = None
                self.combo.set('')
                
                # Update button states
                if hasattr(self, 'add_btn'):
                    self.add_btn.config(state='disabled')
                if hasattr(self, 'add_to_tree_btn'):
                    self.add_to_tree_btn.config(state='disabled')
        else:
            messagebox.showinfo("Info", "No details to clear")

    def show_details(self, product):
        """Show product details in horizontal layout with checkbox"""
        frame = ttk.Frame(self.detail_frame, borderwidth=1, relief="solid", padding=5)
        frame.pack(fill='x', pady=2, padx=7)

        frame.product = product

        # Checkbox variable
        checkbox_var = tk.BooleanVar()
        frame.checkbox_var = checkbox_var

        # Checkbox with callback
        checkbox = ttk.Checkbutton(
            frame,
            variable=checkbox_var,
            command=self.update_add_btn_state  # <-- enable/disable "+"
        )
        checkbox.pack(side='left', padx=(0, 5))

        # Part Number
        part_number = str(product['Part Number'])
        if len(part_number) > 15:
            part_number = part_number[:12] + "..."
        part_label = ttk.Label(frame, text=part_number, font=('Arial', 9, 'bold'), width=15)
        part_label.pack(side='left', padx=(0, 2))

        # Description
        description = str(product['Description'])
        if len(description) > 50:
            description = description[:43] + "..."
        desc_label = ttk.Label(frame, text=description, font=('Arial', 9), width=20)
        desc_label.pack(side='left', padx=(0, 5))

        # Clear button with state update
        clear_btn = ttk.Button(
            frame,
            text="✕",
            width=3,
            command=lambda f=frame: (f.destroy(), self.update_add_btn_state())
        )
        clear_btn.pack(side='right', padx=(0, 0))

    def add_selected_items(self):
        """Add all checked items to the tree"""
        added_count = 0
        frames_to_remove = []
        
        # Check all frames in detail_frame for selected checkboxes
        for child in self.detail_frame.winfo_children():
            if hasattr(child, 'checkbox_var') and child.checkbox_var.get():
                # This frame has a checked checkbox
                product = child.product
                try:
                    # Get formatted values
                    values = self.format_item_for_tree(product)
                    
                    # Insert into tree
                    self.item_tree.insert("", "end", values=values)
                    self.selected_items.append(product)
                    added_count += 1
                    frames_to_remove.append(child)
                    
                except ValueError as e:
                    messagebox.showerror("Error", f"Invalid data format for {product.get('Part Number', 'Unknown')}: {e}")
        
        # Remove the frames of added items
        for frame in frames_to_remove:
            frame.destroy()

        self.update_add_btn_state()
        
        if added_count > 0:
            self.update_remove_button_state()
            self.combo.set('')
            messagebox.showinfo("Items Added", f"Added {added_count} item(s) to the list.")
        else:
            messagebox.showwarning("No Selection", "Please select items using checkboxes first.")

    def update_add_btn_state(self):
        """Enable '+' if any search item is checked"""
        has_checked = any(
            getattr(child, 'checkbox_var', None) and child.checkbox_var.get()
            for child in self.detail_frame.winfo_children()
        )
        if hasattr(self, 'add_btn'):  # make sure button exists
            self.add_btn.config(state='normal' if has_checked else 'disabled')

    def format_item_for_tree(self, product):
        """Format product data for tree display - override in subclasses"""

        qty = format_qty(product.get('Qty', 1))
        price = format_price(product.get('Unit Price', 0), self.currency_unit.get())
        weight = format_weight(product.get('Weight', 0))
        
        return (
            product['Part Number'],
            product['Description'],
            qty,
            product.get('Supplier', ''),
            price,
            weight
        )
    
    def remove_selected_item(self):
        """Remove selected items from the tree"""
        all_items = self.item_tree.get_children()
        num_items = len(all_items)

        if num_items == 0:
            messagebox.showwarning("No Items", "There are no items to remove.")
            self.remove_btn.config(state='disabled')
            return

        if num_items == 1:
            # Only one item - remove it immediately
            self.item_tree.delete(all_items[0])
            if self.selected_items:
                self.selected_items.pop(0)
            self.update_remove_button_state()
            return

        # Multiple items - require selection
        selected = self.item_tree.selection()
        if not selected:
            messagebox.showwarning("No Selection", "Please select item(s) to remove.")
            return

        # Remove selected items
        indices_to_remove = sorted([self.item_tree.index(item) for item in selected], reverse=True)

        for item_id in selected:
            self.item_tree.delete(item_id)

        # Remove from selected_items list
        for index in indices_to_remove:
            if index < len(self.selected_items):
                self.selected_items.pop(index)

        self.update_remove_button_state()

    def select_all_items(self):
        """Check all items currently displayed in the detail frame and add them to treeview"""
        any_item = False
        for child in self.detail_frame.winfo_children():
            if hasattr(child, 'checkbox_var'):
                child.checkbox_var.set(True)
                any_item = True
        
        if any_item:
            self.update_add_btn_state()
            self.add_selected_items()
        else:
            messagebox.showinfo("No Items", "There are no items to select.")

    
    def reset_items(self):
        """Reset all items"""
        self.selected_items.clear()
        for row in self.item_tree.get_children():
            self.item_tree.delete(row)
        self.update_remove_button_state()

    def refresh_treeview(self):
        """Clear and reload the treeview based on current selected_items"""
        # Clear all items in the treeview
        for row in self.item_tree.get_children():
            self.item_tree.delete(row)
        
        # Reinsert items from selected_items
        for product in self.selected_items:
            try:
                values = self.format_item_for_tree(product)
                self.item_tree.insert("", "end", values=values)
            except Exception as e:
                print(f"Error refreshing item {product.get('Part Number', '')}: {e}")
        
        # Update remove button state
        self.update_remove_button_state()

    
    def update_remove_button_state(self):
        """Update remove button state based on items"""
        if self.item_tree.get_children():
            self.remove_btn.config(state='normal')
        else:
            self.remove_btn.config(state='disabled')
    
    def export_to_excel(self):
        """Export items to Excel using get_export_data() with column names"""
        if not self.item_tree.get_children():
            messagebox.showinfo("No Data", "No items to export.")
            return

        try:
            data = self.get_export_data()
            if not data:
                messagebox.showinfo("No Valid Data", "No valid rows to export.")
                return

            default_name = getattr(self, 'default_filename', 'export.xlsx')
            file_path = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel Files", "*.xlsx")],
                title=f"Save {getattr(self, 'module_title', 'Export')} As",
                initialfile=default_name
            )
            if not file_path:
                return

            # Check overwrite
            if os.path.exists(file_path):
                confirm = messagebox.askyesno(
                    "Confirm Overwrite",
                    f"The file '{os.path.basename(file_path)}' already exists.\nDo you want to overwrite it?"
                )
                if not confirm:
                    return

            # Convert list of dicts to DataFrame
            df_new = pd.DataFrame(data)

            # Ensure consistent column order as in treeview
            tree_columns = list(self.item_tree["columns"])
            df_new = df_new.reindex(columns=tree_columns)

            # Export to Excel
            df_new.to_excel(file_path, index=False, engine='openpyxl')

            # Open workbook to set specific column widths
            wb = load_workbook(file_path)
            ws = wb.active

            # Optional: set widths for known columns
            width_map = {
                "Part Number": 20,
                "Description": 40,
                "Qty": 10,
                "Supplier": 20,
                "Unit Price": 15,
                "Weight": 12
            }

            for i, col in enumerate(df_new.columns, 1):
                if col in width_map:
                    ws.column_dimensions[get_column_letter(i)].width = width_map[col]

            wb.save(file_path)
            messagebox.showinfo("Export Successful", f"File saved to:\n{file_path}")

        except Exception as e:
            messagebox.showerror("Export Failed", f"Could not save file:\n{e}")

    def upload_file(self):
        file_path = filedialog.askopenfilename(filetypes=[
            ("Word Documents", "*.docx"),
            ("Excel files", "*.xlsx *.xls")
            ])
        
        if not file_path:
            return None
        
        data = {"items": [], "info": {}}
        if file_path.endswith(('.xlsx', '.xls')):
                df = pd.read_excel(file_path)
                self.products = df.to_dict(orient='records')
                
                # Save to cache
                if save_products_cache(self.products):
                    print("Products cached successfully")
                
                # Update display list
                self.build_combo_display_list()
                self.combo['values'] = self.combo_display_list
                messagebox.showinfo("Success", "Product list uploaded and cached successfully.")
                return data
        elif file_path.endswith('.docx'):
            doc = Document(file_path)
            data = {"info": {}, "items": []}
            if len(doc.tables) < 2:
                messagebox.showerror("Error", "Word template must contain at least 2 tables (info + items).")
                return None

            # ----------------
            # Parse Info Table
            # ----------------
            info_table = doc.tables[0]
            try:
                data["info"]["Customer"] = info_table.cell(0,0).text.strip()
                data["info"]["Delivery_Note_No"] = info_table.cell(0,1).text.strip()
                data["info"]["Project_Name"] = info_table.cell(1,0).text.strip()
                data["info"]["Address"] = info_table.cell(2,0).text.strip()

                data["info"]["Phone_Num"] = info_table.cell(3,0).text.strip()
                data["info"]["Date"] = info_table.cell(3,1).text.replace("Date:", "").strip()

                data["info"]["Incharge"] = info_table.cell(4,0).text.replace("Attn.:", "").strip()
                data["info"]["Contact_Num"] = info_table.cell(5,0).text.replace("Mob", "").strip()
                data["info"]["Customer_PO"] = info_table.cell(6,0).text.replace("Customer Po Ref:", "").strip()
                data["info"]["Quotation"] = info_table.cell(7,0).text.replace("Our Quotation:", "").strip()
                data["info"]["Subject"] = info_table.cell(9,0).text.replace("Subject:", "").strip()
            except IndexError:
                messagebox.showerror("Error", "Word info table format does not match expected template.")
                return None

            # ----------------
            # Parse Items Table
            # ----------------
            items_table = doc.tables[1]
            for row in items_table.rows[1:]:  # skip header
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 4:
                    item_code = cells[1]     # skip "No."
                    desc = cells[2]
                    qty = cells[3]
                    data["items"].append([item_code, desc, qty])

            return data
        
        return data
    
    def on_double_click(self, event):
        """Handle double-click on treeview item to edit"""
        item = self.item_tree.identify_row(event.y)
        column = self.item_tree.identify_column(event.x)

        if not item:
            return
            
        column_id = int(column[1]) - 1
        column_name = self.item_tree["columns"][column_id]
        
        # Only allow editing Qty and unit price columns by default
        if column_name not in ["Qty", "Unit Price", "Weight"]:
            return
            
        current_value = self.item_tree.item(item)['values'][column_id]
        
        # Create entry widget for editing
        entry = ttk.Entry(self.item_tree)
        entry.insert(0, current_value)
        
        bbox = self.item_tree.bbox(item, column)
        entry.place(x=bbox[0], y=bbox[1], width=bbox[2], height=bbox[3])
        
        def on_enter(e):
            new_value = entry.get().strip()
            try:
                if column_name == "Qty":
                    qty_str = ''.join(filter(str.isdigit, new_value))
                    if not qty_str:
                        messagebox.showerror("Invalid Input", "Quantity must be a positive number")
                        return
                    
                    qty_val = int(qty_str)
                    if qty_val <= 0:
                        messagebox.showerror("Invalid Input", "Quantity must be greater than 0")
                        return
                    
                    new_value = format_qty(qty_val)
                
                values = list(self.item_tree.item(item)['values'])
                values[column_id] = new_value
                self.item_tree.item(item, values=values)
                
            except ValueError:
                messagebox.showerror("Invalid Input", f"Please enter a valid {column_name.lower()}")
            finally:
                entry.destroy()
        
        entry.bind('<Return>', on_enter)
        entry.bind('<Escape>', lambda e: entry.destroy())
        entry.focus_set()
    
    def on_enter_pressed(self, event):
        """Handle Enter key in search combobox"""
        self.combo.event_generate('<Button-1>')
    
    def return_home(self):
        """Return to homepage"""
        self.destroy()
        self.parent.deiconify()
    
    def get_live_rate(self, from_currency, to_currency):
        if not API_KEY:
            print("Error: Exchange Rate API key not set.")
            return None
        url = f"https://v6.exchangerate-api.com/v6/{API_KEY}/latest/{from_currency}"
        
        try:
            response = requests.get(url)
            data = response.json()
            
            if data["result"] != "success":
                print("Error fetching exchange rate:", data.get("error-type"))
                return None
            
            rate = data["conversion_rates"].get(to_currency)
            if rate is None:
                print(f"Currency {to_currency} not found.")
                return None
            
            return rate
        except Exception as e:
            print("Error fetching exchange rate:", e)
    
    def on_currency_changed(self, event=None):
        target_currency = self.currency_var.get()
        self.currency_unit.set(target_currency)

        rate = 1
        if target_currency != "SAR":
            rate = self.get_live_rate("SAR", target_currency)
            if not rate:
                rate = 1  # fallback

        for item_id in self.item_tree.get_children():
            values = list(self.item_tree.item(item_id)['values'])
            price_str = str(values[4])
            # Remove currency symbol and commas
            for symbol in ["SAR", "USD", ","]:
                price_str = price_str.replace(symbol, "")
            try:
                sar_price = float(price_str.strip())
                new_price = sar_price * rate
                # Use your format_price here to maintain consistency
                values[4] = format_price(new_price, target_currency)
                self.item_tree.item(item_id, values=values)
            except ValueError:
                continue



    # Canvas scroll methods
    def on_frame_configure(self, event=None):
        """Reset the scroll region to encompass the inner frame"""
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))

    def on_canvas_configure(self, event):
        """When canvas is resized, resize the inner frame to match"""
        width = event.width
        self.canvas.itemconfig(self.canvas_frame, width=width)

    def on_main_frame_configure(self, event=None):
        """Reset the scroll region to encompass the main frame"""
        self.main_canvas.configure(scrollregion=self.main_canvas.bbox("all"))

    def on_main_canvas_configure(self, event):
        """When main canvas is resized, resize the main frame to match"""
        width = event.width
        self.main_canvas.itemconfig(self.main_canvas_frame, width=width)